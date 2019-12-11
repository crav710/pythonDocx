from docx import Document
import pandas as pd
import os 
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import time 
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches





def iter_block_items(parent):
	if isinstance(parent, _Document):
		parent_elm = parent.element.body
	elif isinstance(parent, _Cell):
		parent_elm = parent._tc
	elif isinstance(parent, _Row):
		parent_elm = parent._tr
	else:
		raise ValueError("something's not right")
	for child in parent_elm.iterchildren():
		if isinstance(child, CT_P):
			yield Paragraph(child, parent)
		elif isinstance(child, CT_Tbl):
			yield Table(child, parent)


def generatefiletags(input_wordfile,tags,tags_dict):
	document=Document(input_wordfile)
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			# print('bt ',block_text)
			for i in range(len(tags)):
				tag=tags[i]
				if tag in block_text:
					complete_tag=tag+(block_text.split(tag)[1].split(']]')[0])+']]'
					# print('cTag: ',complete_tag)
					if i in tags_dict:
						tags_dict[i].append(complete_tag)
					else:
						tags_dict[i]=[complete_tag]
		elif isinstance(block, Table):
			for row in block.rows:
				row_data = []
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						block_text=paragraph.text
						for i in range(len(tags)):
							tag=tags[i]
							if tag in block_text:
								complete_tag=tag+(block_text.split(tag)[1].split(']]')[0])+']]'
								if i in tags_dict:
									tags_dict[i].append(complete_tag)
								else:
									tags_dict[i]=[complete_tag]						
						row_data.append(paragraph.text)
				# print("\t".join(row_data))
	# print(tags_dict)
	return tags_dict



def generate_tags_list(input_wordfile,list_df):
	tags=['[[FILE:','[[IMAGE:','[[TEXT:','[[IF:']
	tags_dict={}
	tags_dict=generatefiletags(input_wordfile,tags,tags_dict)




def getallfiles(list_df):
	file_df=list_df[3]
	filecolumns=list(file_df.columns)[1:]
	# print(filecolumns)
	total_files=[]
	for col in filecolumns:
		total_files=total_files+list(file_df[col])
	# print(total_files)
	total_files=list(dict.fromkeys(total_files))
	# print(total_files)
	return total_files

def evaluateTag(list_df,tag,i):

	glob_df=list_df[0]
	tag_value=''
	try:

		tag_value=list(glob_df[(tag.split(':')[1]).split(']]')[0]])[i]
		if tag_value==True:
			tag_value='TRUE'
		else:
			tag_value='FALSE'
	except Exception as e:
		print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))


	return tag_value


def replaceIftext(list_df,block_text,i):
	replace_text=""
	if_count=0
	else_count=0
	endif_count=0
	waitforendif=0
	waitforelse=0
	leave_tags_detected=0
	if_tags=[]
	if_tags_order=[]
	else_tags_order=[]
	text_replace=[]
	other_tags=['[[ELSE]]','[[ENDIF]]']
	leave_tags=['[[TEXT:','[[IMAGE:']
	split_data=block_text.split('[[IF:')
	# print(split_data)
	for x in split_data[1:]:
		if_tags.append('[[IF:'+x.split(']]')[0]+']]')
	# print(if_tags)
	dtect_tag=''
	leave_tag_text=''
	start=0
	tag_text=''
	z='TRUE'
	# print(block_text)
	for x in block_text:
		# print(x)
		if x=='[':
			start=1
		if start==1:
			dtect_tag=dtect_tag+x
			
		if start==0:
			tag_text=tag_text+x

		if dtect_tag in if_tags:
			if waitforendif==1:
				start=0
				dtect_tag=''
			elif waitforelse==1:
				start=0
				dtect_tag=''
			else:
				# print('Tag_detected: ',dtect_tag)
				# print('Tag Text: ',tag_text)
				if z =='FALSE':
					waitforelse=1
				else:
					text_replace.append(tag_text)
				z=evaluateTag(list_df,dtect_tag,i)
				# print('eval tag : ',z)
				if_tags_order.append(z)
				else_count=else_count+1
				# print('text Replace :',text_replace)
				# print('if_tags_order:',if_tags_order)			
				tag_text=''
				dtect_tag=''
				start=0
		if dtect_tag in other_tags:
			# print('Tag_detected other: ',dtect_tag)
			# print('Tag Text: ',tag_text)
			if waitforendif==1:
				if dtect_tag==other_tags[1]:
					# print('ENDIF')
					start=0
					dtect_tag=''
					tag_text=''
					waitforendif=0
			else:
				if z=='TRUE':
					text_replace.append(tag_text)
				# print('text Replace :',text_replace)
				if dtect_tag==other_tags[0]:
					tag_eval=if_tags_order[else_count-1]
					else_count=else_count-1
					# print('tag Eval :',tag_eval)
					if tag_eval=='TRUE':
						waitforendif=1
						start=0
						dtect_tag=''
						tag_text=''
					else:
						z='TRUE'
						dtect_tag=''
						tag_text=''
						start=0
				else:
					dtect_tag=''
					tag_text=''
					start=0
					z='TRUE'

		# print(dtect_tag)
		# print(other_tags)
		if dtect_tag in leave_tags:
			# print('dtect_tag extra:',dtect_tag)
			tag_text=tag_text+dtect_tag
			start =0
			dtect_tag=''
	# print(text_replace)
	return (''.join(text_replace))


def find_and_replaceiftags(list_df,file,i):
	# tags_associated=['[[ELSE]]',]
	document=Document(file)
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			if '[[IF:' in block_text:
				# print('btt',block_text)
				z=replaceIftext(list_df,block_text,i)
				block.text=z
			# print('bt ',block_text)
	document.save(file)


def Document_data(document):
	# Run_list=[]
	content=''
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			# print('bt ',block_text)
			# for run in block.runs:
				# print('Run blocks: ',run.text)
			# Run_list.append
			content=content+block_text
	return content

def replaceFiletags(list_df,document,files_dict,i):
	instance=False
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			# print('bt ',block_text)
			file_df=list_df[3]
			if '[[FILE:' in block_text:
				numfiles=block_text.split('[[FILE:')
				# print(len(numfiles))
				for z in numfiles[1:]:
					identifier=z.split(']]')[0]
					# print('idd',identifier)
					tag='[[FILE:'+identifier+']]'
					# print('filename',file_df[identifier])
					try:
						file=list(file_df[identifier])[i]
						# print('content: ',files_dict[file])
						# print('tag',tag)
						block_text=block_text.replace(tag,files_dict[file]+' ')
					except Exception as e:
						print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))
						
					# print('block_text',block_text)
					# block_text=block.text
				block.text=block_text
				instance=True
				# print('Replaced: ',block.text)
		elif isinstance(block, Table):
			# print('Enter Table: ',table)
			# ll
			for row in block.rows:
				row_data = []
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						block_text=paragraph.text
						# print('bt ',block_text)
						file_df=list_df[3]
						if '[[FILE:' in block_text:
							numfiles=block_text.split('[[FILE:')
							# print(len(numfiles))
							for z in numfiles[1:]:
								identifier=z.split(']]')[0]
								# print('idd',identifier)
								tag='[[FILE:'+identifier+']]'
								# print('filename',file_df[identifier])
								try:
									file=list(file_df[identifier])[i]
									# print('content: ',files_dict[file])
									# print('tag',tag)
									block_text=block_text.replace(tag,files_dict[file]+' ')
								except Exception as e:
									print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))
								# print('block_text',block_text)	
								# block_text=paragraph.text
							paragraph.text=block_text
							instance=True
							# print('Replaced: ',paragraph.text)
	return document,instance








	
def load_file_dict(list_df,files):
	files_dict={}
	for file in files:
		document=Document(file)
		files_dict[file]=Document_data(document)
	# print(files_dict)
	return files_dict

def preprocess_files(input_wordfile,list_df,i):
	files=getallfiles(list_df)
	files_dict=load_file_dict(list_df,files)
	files.append(input_wordfile)
	for i in range(len(files)):
		file=files[i]

		if i==len(files)-1:
			files_dict=load_file_dict(list_df,files)
			# ll
		find_and_replaceiftags(list_df,file,i)
		# ll
		document=Document(file)
		# content=Document_data(document)
		while True:			
			document,contain=replaceFiletags(list_df,document,files_dict,i)
			# print('contain: ',contain)
			if contain==True:
				pass
			else:
				break	

		document.save(file)









def validatesheetnames(sheet_names):
	if sheet_names[0]=='GLOBALS' and sheet_names[1]=='TEXT' and sheet_names[2]=='IMAGE' and sheet_names[3]=='FILE':
		return 1
	return 0

def replacehead(df):
	new_header = df.iloc[0] #grab the first row for the header
	df = df[1:] #take the data less the header row
	df.columns = new_header #set the header row as the df header
	return df


def cleanExcel(df):
	df_globals=replacehead(df.parse('GLOBALS').dropna(how='all').dropna(axis='columns'))
	df_text=replacehead(df.parse('TEXT').dropna(how='all').dropna(axis='columns'))
	df_image=replacehead(df.parse('IMAGE').dropna(how='all').dropna(axis='columns'))
	df_file=replacehead(df.parse('FILE').dropna(how='all').dropna(axis='columns'))
	# print(df_file)
	return df_globals,df_text,df_image,df_file

def read_excel(file_name):
	df=pd.ExcelFile(file_name)
	# print(df.sheet_names)
	if validatesheetnames(df.sheet_names):
		df_globals,df_text,df_image,df_file=cleanExcel(df)
		# print(df_globals)
		# print(df_text)
		# print(df_image)
		# print(df_file)
		return [df_globals,df_text,df_image,df_file]
	else:
		return 0	

def replace_text_tags(list_df,input_wordfile,target_file,i):
	document=Document(input_wordfile)
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			# print(block_text)
			if '[[TEXT:' in block_text:
				numtexttags=block_text.split('[[TEXT:')
				# print('numtexttags: ',numtexttags)
				for tag in numtexttags[1:]:
					identifier=tag.split(']]')[0]
					# print('IDD: ',identifier)
					tag_d='[[TEXT:'+identifier+']]'
					# print('Tag ID : ',tag_d)
					try:
						text_rep=list(list_df[1][identifier])[i]
						# print('text_rep: ',text_rep)
						block.text=block_text.replace(tag_d,text_rep+' ')
						# print('bt ',block.text)
						block_text=block.text
					except Exception as e:
						print('Error the Tag {} is not present in the Excel Sheet.'.format(tag_d))

		elif isinstance(block, Table):
			for row in block.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						block_text=paragraph.text
						if '[[TEXT:' in block_text:
							numtexttags=block_text.split('[[TEXT:')
							# print('numtexttags: ',numtexttags)
							for tag in numtexttags[1:]:
								identifier=tag.split(']]')[0]
								# print('IDD: ',identifier)
								tag_d='[[TEXT:'+identifier+']]'
								# print('Tag ID : ',tag_d)
								try:
									text_rep=list(list_df[1][identifier])[i]
									# print('text_rep: ',text_rep)
									block_text=block_text.replace(tag_d,text_rep+' ')
									# print('block_text: ',block_text)
								except Exception as e:
									print('Error the Tag {} is not present in the Excel Sheet.'.format(tag_d))
							paragraph.text=block_text
							# print('bt ',paragraph.text)
	document.save(target_file)

def insert_image_after(run, image,img_width,img_height):
    """Insert a new paragraph after the given paragraph."""
    run.add_picture(image,width = Inches(img_width), height = Inches(img_height))
    
def insert_run_after(paragraph, text):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para
def createParagraph(paragraph,list_images):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    insert_image_after(new_para.add_run())

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def replace_image_tag(list_df,input_wordfile,target_file,i):
	document=Document(input_wordfile)
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			
			if '[[IMAGE:' in block.text:
				for run in block.runs:
					# print('run : ',run.text)
					try:
						if '[[IMAGE:' in run.text:
							image_items=(run.text).split('[[IMAGE:')
							# print('image_items: ',image_items)
							new_p = OxmlElement("w:p")
							block._p.addnext(new_p)
							new_para = Paragraph(new_p, block._parent)
							new_para.add_run(image_items[0])				
							for image in image_items[1:]:
								# print('split image : ',image.split(']]'))
								image_tag='[[IMAGE:'+image.split(']]')[0]+']]'
								idd=image.split(']]')[0]
								text_add=image.replace((idd+']]'),' ')
								# print('text_add: ' ,text_add)
								try:
									img_path=list(list_df[2][idd])[i]
									img_width=list(list_df[2][idd+'_width'])[i]
									img_height=list(list_df[2][idd+'_height'])[i]
									insert_image_after(new_para.add_run(),img_path,img_width,img_height)
									new_para.add_run(text_add)
								except Exception as e:
									new_para.add_run(text_add)
									# block.text=(block.text).replace(text_add,'')
									# print('ErrorPar: ',str(e))
									print('Error the Tag {} is not present in the Excel Sheet.'.format(image_tag))	
					except Exception as e:
						# print('ErrorPar: ',str(e))
						print('Image Tag does not fall under  Run.')
						# run.text=''
				delete_paragraph(block)
			# print('bt ',block_text)
		elif isinstance(block, Table):
			for row in block.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						block_text=paragraph.text
						if '[[IMAGE:' in paragraph.text:
							for run in paragraph.runs:
								# print('run: ',run.text)
								try:
									if '[[IMAGE:' in run.text:
										image_items=(run.text).split('[[IMAGE:')
										# print('image_items: ',image_items)
										new_p = OxmlElement("w:p")
										paragraph._p.addnext(new_p)
										new_para = Paragraph(new_p, paragraph._parent)
										new_para.add_run(image_items[0])
										for image in image_items[1:]:
											# print('split image : ',image.split(']]'))
											image_tag='[[IMAGE:'+image.split(']]')[0]+']]'
											idd=image.split(']]')[0]
											text_add=image.replace((idd+']]'),' ')
											# print('text_add: ' ,text_add)
											try:
												img_path=list(list_df[2][idd])[i]
												img_width=list(list_df[2][idd+'_width'])[i]
												img_height=list(list_df[2][idd+'_height'])[i]
												insert_image_after(new_para.add_run(),img_path,img_width,img_height)
												new_para.add_run(text_add)
											except Exception as e:
												new_para.add_run(text_add)
												# block.text=(block.text).replace(text_add,'')
												# print('ErrorPar: ',str(e))
												print('Error the Tag {} is not present in the Excel Sheet.'.format(image_tag))	
								except Exception as e:							
									# print('ErrorTable: ',str(e))
									print('Image Tag Doesnt fall under run ')
							delete_paragraph(paragraph)

	document.save(target_file)






def read_document(list_df,input_wordfile,target_file,i):
	find_and_replaceiftags(list_df,input_wordfile,i)
	replace_text_tags(list_df,input_wordfile,target_file,i)
	replace_image_tag(list_df,target_file,target_file,i)













if __name__ == '__main__':
	print('Loading And Validating EXCEL')
	excel_file_path=os.path.join(os.path.join(os.getcwd(),'Files'),'excelfile.xlsx')
	list_df=read_excel(excel_file_path)
	input_wordfile=os.path.join(os.path.join(os.getcwd(),'Files'),'inputDocument.docx')
	# generate_tags_list(input_wordfile,list_df)
	target_file=list(list_df[0]['DESTINATION'])
	# print(target_file)
	# ll

	for i in range(len(target_file)):
		print('Preprocessing Files')
		preprocess_files(input_wordfile,list_df,i)
		print('Iteration ',i+1)
		file_target=target_file[i]
		read_document(list_df,input_wordfile,file_target,i)

	# read_document(list_df,input_wordfile,target_file[0],0)