from docx import Document
import pandas
import os
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
import time
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches
import xlrd  # only listed here to ensure that the Pandas dependency is met


# 12/18/19 comments
# - getalltags() function was renamed to identify_unique_filepaths_from_all_file_tags()
#   to better desribe what the function did
# - Document_data() function was renamed to identify_unformatted_runs_from_document()
#   to better describe what the function did
#
# Dave's most important comments
#
# I understand that [[IF...]] tag processing is still in progress so I did
# not look over that code
#
# Tag replacement - including [[FILE...]] tag replacement - seems to lose
# all formatting information (except for [[IMAGE...]] tag replacement as
# discussed below) because tag replacement relies on extracting unformatted
# content for a Paragraph, modifying that content, and storing the modified
# content into the Paragraph instance.  Thus, any formatting within the
# Paragraph will be lost.  This general approach should work, without
# losing formatting, if this operation is done on a Run level rather than
# a Paragraph level, because formatting is unchanged within a Run.
# However, [[IMAGE...]] tag replacement operates at a Run level and seems
# to avoid this issue, thereby suggesting that if Ravi confirms the issue
# with formatting loss, a solution is already shown for the issue
#
# Tag parsing does not consider the presence of optional comments surrounded
# by parentheses; e.g., [[FILE=identifer99 (this is the main file)]].  This
# feature is not high priority and can be deferred for now (possibly to a
# follow-on project that would expand the current project)
#
# I understand that Excel file processing and validation is in progress
#
# Note that I commented out logic for [[IF...]] tag processing and did
# not review that code because I understand that nested [[IF...]] tag
# processing is in progress
#
# The functions below are not tied into the executing code at present, so I
# have not reviewed them (with the caveat that I chose not to review code
# for [[IF...]] tag processing because I understand that logic for handling
# nested [[IF...]] tags is in progress.
# - generatefiletags
# - generate_tags_list
# - evaluateTag
# - replaceIftext


def iter_block_items(parent):  # CONFIRMED
    """
    Provides an iterator interface for retrieving python-docx Paragraph and
    Table entities from a parent/containing entity
    :param parent: (python.docx instance of some type) parent entity
    :return: (iterator output) emits a series of python-docx.Paragraph or Table
    instances based on the content of the parent entity
    """

    # Performs type-specific processing, based on data type of parent, to
    # extract parent-data-type-specific data from the parent such as paragraphs
    # or tables from a Document.  This step sets up an easy capability to iterate
    # through the children of parent on a parent-type-independent basis
    if isinstance(parent, _Document):
        # parent is a python-docx Document so set parent element to the
        # document's body container
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        # parent is a python-docx table cell so set parent element to the container
        # for the table cell
        parent_element = parent._tc
    elif isinstance(parent, _Row):
        # parent is a python-docx table row so set parent element to the container
        # for the table row
        parent_element = parent._tr
    else:
        # parent is none of document, table cell or table row so error
        raise ValueError("something's not right")

    # Iterate through each child element within the parent element's container
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            # current child is a paragraph so emit a python-docx.Paragraph instance
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            # current child is a table so emit a python-docx.Table instance
            yield Table(child, parent)


#  complete verification later.


def identify_unique_filepaths_from_all_file_tags(ordered_dataframes):  # Confirmed3
    """
    Determine a list of unique filepaths, for all FILE tags, for all
    iterator values in the FILE worksheet.  This function reflects
    an optimization to parse Word files at these filepaths once
    up front rather than parsing those files potentially many times
    :param ordered_dataframes: (list of Pandas dataframes) Pandas dataframes
    containing workbook data
    :return: (list of strings) a unique list of all filepaths for
    all FILE tags and iterator values
    """

    # Extract FILES (the 3rd indexed) data frame from data frames
    # read from workbook
    # Action Item: convert 3 to a oonstant defined up front
    file_dataframe = ordered_dataframes[3]

    # Retrieve columns from FILES worksheet (not including 0th
    # column, which is the iterator column).  Note that column header
    # row was previously removed so we only get data from this call
    columns_of_filepaths = list(file_dataframe.columns)[1:]

    # print(filecolumns)

    # Create an empty list for storing FILE filepaths prior to
    # appending individual results to that list
    filepaths = []

    # Iterate through each column in the FILE workbook, adding
    # files for that column to the running list of files, so that
    # we can develop a list of all filepaths for all iterator values
    for filepath_column in columns_of_filepaths:
        filepaths = filepaths + list(file_dataframe[filepath_column])

    # print(total_files)

    # Remove duplicates in the list by converting the
    # list to a dictionary, with keys matching values for
    # each entry, and then converting the resulting
    # dictionary back into a list
    unique_filepaths = list(dict.fromkeys(filepaths))

    # print(total_files)
    return unique_filepaths


def evaluateTag(list_df, tag, i):  # CONFIRMED
    """
    Determines whether the identifier for the provided tag is
    defined in the GLOBALS worksheet for the current iterator
    and returns "TRUE" if it is otherwise returns "FALSE"
    :param list_df: (list of Pandas data frames) workbook data
    :param tag: (string)
    :param i: (int) iterator value
    :return: (string) "TRUE" or "FALSE" reflecting whether the
    provided tag is defined in the GLOBALS worksheet for the
    current iterator.  "FALSE" represents an error
    """

    # Retrieve GLOBAL worksheet data from data frames
    # Action Item: convert 0 to global variable defined above
    glob_df = list_df[0]

    # Extract the identifier and comment, if any, for the tag
    tag_data = (tag.split(':')[1]).split(']]')[0]

    # Check whether comment is present
    if '(' in tag_data:
        # Extract identifier from tag excluding spaces
        tag_data = tag_data.split('(')[0].strip()

    # Check whether tag's identifier is defined within GLOBAL
    # sheet for current iterator
    tag_value = list(glob_df[tag_data])[i]

    # Convert Boolean True to "TRUE" else "FALSE"
    if tag_value == True:
        tag_value = 'TRUE'
    else:
        tag_value = 'FALSE'

    # Return "TRUE" or "FALSE"
    return tag_value




def replace_if_else(listdf, block, index):
    #  This function will run x number of times. Where x is the nested level of tags.
    runlist=block.runs
    while True:
        if_count = 0
        else_count = 0
        temp_list=[]
        runlist = temp_list
        tag_eval = 'TRUE'
        ignore_text = 0
        for run in block.runs:
            # print('run : ',run.text)
            try:
                if '[[IF:' in run.text:
                    substring = (run.text).split('[[IF:')
                    identifier = substring.split(']]')[0]
                    if '(' in identifier:
                        identifier = identifier.split('(')[0]
                        comment = identifier.split('(')[1].split(')')[0]
                        tag = '[[IF:' + identifier + '(' + comment + ')' + ']]'
                    else:
                        tag = '[[IF:' + identifier + ']]'
                    if tag==run.text:
                        if_count = if_count + 1
                        if if_count == 1:
                            tag_eval = evaluateTag(listdf, tag, index)
                            if tag_eval == 'FALSE':
                                ignore_text = 1
                        if tag_eval == 'TRUE':
                            if if_count > 1 and ignore_text == 0:
                                temp_list.append(run)
                    else:
                        print('Tag {} does not completely lie in a run.')
                elif '[[ELSE]]' in run.text:
                    else_count = else_count + 1
                    if else_count != if_count:
                        if tag_eval == 'TRUE' and ignore_text == 0:
                            temp_list.append(run)
                    else:
                        if tag_eval == 'FALSE' and ignore_text == 0:
                            temp_list.append(run)
                        else:
                            ignore_text = 1
                else:
                    if tag_eval == 'TRUE' and ignore_text == 0:
                        temp_list.append(run)
            except:
                print('IF Tag no Present')
        if if_count == 0:
            break
    return runlist

def replace_nested_tags(listdf, filename, index):
    document = Document(filename)
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            block_text = block.text
            if '[[IF:' in block.text:
                new_p = OxmlElement("w:p")
                block._p.addnext(new_p)
                new_para = Paragraph(new_p, block._parent)
                run_list = replace_if_else(listdf, block, index)
                for runs in run_list:
                    sentence=new_p.add_run(runs.text)
                    setAttributes(sentence,runs)
                delete_paragraph(block)
        # print('bt ',block_text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        block_text = paragraph.text
                        if '[[IF:' in paragraph.text:
                            new_p = OxmlElement("w:p")
                            block._p.addnext(new_p)
                            new_para = Paragraph(new_p, block._parent)
                            run_list = replace_if_else(listdf, block, index)
                            for runs in run_list:
                                new_p.add_run(runs.text)
                            delete_paragraph(paragraph)

    document.save(filename)




def identify_unformatted_runs_from_document(document):  # CONFIRMED2
    """
    Returns a list of strings containing all unformatted text
    from all Paragraph instances in the provided Document instance
    :param document: (python-docx.Document) input Word file
    :return: (string) unformatted content from all Paragraph
    instances within the provided Document instance
    """

    # Create an empty list of informatted strings to which
    # unformatted strings extracted from Run instances will be added
    para_count=0
    para_runs={}
    # Iterate through each Paragraph and Table instance
    # within the provided Document instance
    for block in iter_block_items(document):
        # Only process content for Paragraph instances
        if isinstance(block, Paragraph):
            para_count=para_count+1
            # Retrieve unformatted text from Paragraph
            # (no longer used due to coding change)
            # block_text=block.text

            # print('bt ',block_text)
            run_list=[]

            # Iterate through each Run instance within the Block
            for run in block.runs:
                # Add unformatted text for current Run instance to list
                # of unformatted text instances
                run_list.append(run)
            para_runs[para_count]=run_list
    # Return string containing all unformatted text from
    # block (which can only be a Paragraph)
    return para_runs


def setAttributes(sentence,current_run):
    # sentence=current_run
    # Setting style attributes
    sentence.style.style_id=current_run.style.style_id
    #  Setting Base Style :
    try:
        sentence.style.base_style.base_style = current_run.style.base_style.base_style
        sentence.style.base_style.builtin = current_run.style.base_style.builtin
        sentence.style.base_style.element = current_run.style.base_style.element
        sentence.style.base_style.hidden = current_run.style.base_style.hidden
        sentence.style.base_style.name = current_run.style.base_style.name
        sentence.style.base_style.priority = current_run.style.base_style.priority
        sentence.style.base_style.quick_style = current_run.style.base_style.quick_style
        sentence.style.base_style.unhide_when_used = current_run.style.base_style.unhide_when_used
        sentence.style.base_style.style_id = current_run.style.base_style.style_id

        sentence.style.base_style.font.complex_script = current_run.style.base_style.font.complex_script
        sentence.style.base_style.font.all_caps = current_run.style.base_style.font.all_caps
        sentence.style.base_style.font.name = current_run.style.base_style.font.name
        sentence.style.base_style.font.size = current_run.style.base_style.font.size
        sentence.style.base_style.font.italic = current_run.style.base_style.font.italic
        sentence.style.base_style.font.emboss = current_run.style.base_style.font.emboss
        sentence.style.base_style.font.underline = current_run.style.base_style.font.underline
        sentence.style.base_style.font.highlight_color = current_run.style.base_style.font.highlight_color
        sentence.style.base_style.font.bold = current_run.style.base_style.font.bold
        sentence.style.base_style.font.cs_bold = current_run.style.base_style.font.cs_bold
        sentence.style.base_style.font.cs_italic = current_run.style.base_style.font.cs_italic
        sentence.style.base_style.font.double_strike = current_run.style.base_style.font.double_strike
        sentence.style.base_style.font.hidden = current_run.style.base_style.font.hidden
        sentence.style.base_style.font.imprint = current_run.style.base_style.font.imprint
        sentence.style.base_style.font.math = current_run.style.base_style.font.math
        sentence.style.base_style.font.no_proof = current_run.style.base_style.font.no_proof
        sentence.style.base_style.font.outline = current_run.style.base_style.font.outline
        sentence.style.base_style.font.strike = current_run.style.base_style.font.strike
        sentence.style.base_style.font.superscript = current_run.style.base_style.font.superscript
        sentence.style.base_style.font.subscript = current_run.style.base_style.font.subscript
        sentence.style.base_style.font.web_hidden = current_run.style.base_style.font.web_hidden
        sentence.style.base_style.font.color.rgb = current_run.style.base_style.font.color.rgb
        sentence.style.base_style.font.color.theme_color = current_run.style.base_style.font.color.theme_color

    except:
        print('Base Style Attributes not present.')
    try:
        sentence.style.builtin = current_run.style.builtin
    except:
        print('Builtin attribute not present.')

    # sentence.style.element=current_run.style.element
    sentence.style.hidden=current_run.style.hidden
    sentence.style.name=current_run.style.name
    sentence.style.priority=current_run.style.priority
    sentence.style.quick_style=current_run.style.quick_style
    sentence.style.unhide_when_used=current_run.style.unhide_when_used
    #  Setting Style Font Attributes
    # sentence.style.font=current_run.style.font
    sentence.style.font.complex_script=current_run.style.font.complex_script
    sentence.style.font.all_caps = current_run.style.font.all_caps
    sentence.style.font.name = current_run.style.font.name
    sentence.style.font.size=current_run.style.font.size
    sentence.style.font.italic=current_run.style.font.italic
    sentence.style.font.emboss=current_run.style.font.emboss
    sentence.style.font.underline=current_run.style.font.underline
    sentence.style.font.highlight_color=current_run.style.font.highlight_color
    sentence.style.font.bold=current_run.style.font.bold
    sentence.style.font.cs_bold = current_run.style.font.cs_bold
    sentence.style.font.cs_italic = current_run.style.font.cs_italic
    sentence.style.font.double_strike = current_run.style.font.double_strike
    sentence.style.font.hidden = current_run.style.font.hidden
    sentence.style.font.imprint = current_run.style.font.imprint
    sentence.style.font.math = current_run.style.font.math
    sentence.style.font.no_proof = current_run.style.font.no_proof
    sentence.style.font.outline = current_run.style.font.outline
    sentence.style.font.strike=current_run.style.font.strike
    sentence.style.font.superscript=current_run.style.font.superscript
    sentence.style.font.subscript = current_run.style.font.subscript
    sentence.style.font.web_hidden=current_run.style.font.web_hidden
    sentence.style.font.color.rgb=current_run.style.font.color.rgb
    sentence.style.font.color.theme_color=current_run.style.font.color.theme_color



    # Run high level attributes.

    sentence.element = current_run.element
    sentence.italic = current_run.italic
    sentence.bold = current_run.bold
    sentence.underline=current_run.underline
    # sentence.part=current_run.part



    # Font related attributes
    sentence.font.complex_script=current_run.font.complex_script
    sentence.font.all_caps = current_run.font.all_caps
    sentence.font.name = current_run.font.name
    sentence.font.size=current_run.font.size
    sentence.font.italic=current_run.font.italic
    sentence.font.emboss=current_run.font.emboss
    sentence.font.underline=current_run.font.underline
    sentence.font.highlight_color=current_run.font.highlight_color
    sentence.font.bold=current_run.font.bold
    sentence.font.cs_bold = current_run.font.cs_bold
    sentence.font.cs_italic = current_run.font.cs_italic
    sentence.font.double_strike = current_run.font.double_strike
    sentence.font.hidden = current_run.font.hidden
    sentence.font.imprint = current_run.font.imprint
    sentence.font.math = current_run.font.math
    sentence.font.no_proof = current_run.font.no_proof
    sentence.font.outline = current_run.font.outline
    sentence.font.strike=current_run.font.strike
    sentence.font.superscript=current_run.font.superscript
    sentence.font.subscript = current_run.font.subscript
    sentence.font.web_hidden=current_run.font.web_hidden
    # sentence.font.element=current_run.font.element

#     color attributes
    sentence.font.color.rgb=current_run.font.color.rgb
    sentence.font.color.theme_color=current_run.font.color.theme_color
    # sentence.font.color.type = current_run.font.color.type

def add_paragraph_after(paragraph,runs):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    for run in runs:
        # Add the current unformatted text, as a Run
        # instance, to the new Paragraph instance
        sentence=new_paragraph.add_run(run.text)
        setAttributes(sentence,run)
    return new_paragraph



def replace_file_tags(dataframes,
                      document,
                      filepaths_and_corresponding_lists_of_unformatted_strings,
                      index):  # CONFIRMED3
    """
    Replaces ONE [[FILE...]] tag in python-docx.Document instance with
    content from corresponding Word file.  Caller will repeatedly call
    this function until returned instance is false, reflecting that no
    [[FILE...]] tags were processed, to ensure that all [[FILE...]]
    tags within the Document instance have been processed.  Returns modified
    Document instance and flag indicating whether [[FILE...]] tag was replaced
    :param dataframes: (list of Panda data frames) worksheets
    :param document: (python-docx.Document instance) Word document
    whose FILE tags will be replaced
    :param filepaths_and_corresponding_lists_of_unformatted_strings:
    (dictionary) keys as filepaths and values as a string of unformatted
    content from Paragraphs in the Word document at that filepath
    :param index: (int) 0..n-1 loop index
    :return: document,replacement_done: (python-docx.Document, bool) document
    is the modified Word file after replacing [[FILE...]] tags and
    replacement_done is a flag indicating whether any [[FILE...]] tags were
    replaced
    """

    # replacement_done is True if [[FILE...]] tag replacement has occurred;
    # initialized to false
    replacement_done = False
    pass_replacements=0
    # Iterate through each python-docx.Paragraph or Table instance
    # in received Document instance
    for block in iter_block_items(document):

        # Process as python-docx.Paragraph instance if it is one
        if isinstance(block, Paragraph):

            # Retrieve unformatted text from current block (paragraph)
            block_text = block.text

            # print('bt ',block_text)

            # Retrieve FILES data frame for FILES workbook
            # Action Item: convert sheet index to constant defined above
            file_dataframe = dataframes[3]

            # Check whether start of FILE tag is present within
            # unformatted text for block.  Unformatted text is a
            # viable way to check for the tag because we don't
            # care about formatting considerations (distinguishing
            # Runs) for this check but we care about retaining proper
            # formatting when inserting text
            if '[[FILE:' in block_text:

                # Parse unformatted text into substrings on the
                # [[FILE...]] tag start.  Produces a list of substrings
                # with the 0th element reflecting text preceding the
                # initial [[FILE...]] tag

                # -------------------------------------------------------
                # Add a paragraph after the current paragraph; taken from
                # https://stackoverflow.com/questions/48663788/python-docx-insert-a-paragraph-after/53762537#53762537
                # -------------------------------------------------------

                # create paragraph element
                new_p = OxmlElement("w:p")

                # Append new paragraph to end of current block's container
                block._p.addnext(new_p)

                # Create new Paragraph instance for new paragraph element
                # and link instance to block's parent element
                new_paragraph = Paragraph(new_p, block._parent)
                para = new_paragraph
                # Iterate through runs within block (paragraph)
                for run in block.runs:

                    # See if [[FILE: is in unformatted text of current run.
                    # This check is viable because formatting cannot change,
                    # by design, within a [[FILE...]] tag
                    if '[[FILE:' in run.text:

                        # Yes, [[FILE: was present within unformatted text of run

                        # Split current run into substrings so that text preceding
                        # first [[FILE: text is the first substring, text after the
                        # first [[FILE: text and up to the second [[FILE: (or the
                        # end of the string) is the 2nd substring, etc.
                        substrings = (run.text).split('[[FILE:')

                        # Retrieve the text before the first [[FILE: substring,
                        # create a run for that substring, and add the run to
                        # the paragraph
                        # Question: we seem to lose formatting info here
                        sentence=new_paragraph.add_run(substrings[0])
                        setAttributes(sentence, run)

                        # Iterate through all substrings (excluding the initial
                        # substring because it precedes the first tag)
                        for substring in substrings[1:]:

                            # Extract the [[FILE...]] tag's identifier
                            # (which indexes a filepath in the data frame)
                            # and optional comment from the current substring
                            identifier = substring.split(']]')[0]

                            # Check whether a comment exists
                            if '(' in identifier:
                                # A comment exists

                                # Action Item: this logic assumes that no spaces exist
                                # between the identifier and the comment

                                # Split identifier into substrings on "("
                                # to create an initial substring that is the
                                # identifier with zero or more spaces and the
                                # next substring is the comment through the
                                # closing brace
                                identifier = identifier.split('(')[0]

                                # Extract the text between "(" and ")"
                                comment = identifier.split('(')[1].split(')')[0]

                                print("Processing [[FILE...]] tag with comment (%s)" % comment)

                                # Action Item: this logic assumes that no spaces exist after the
                                # identifier but before the comment, or after the comment but
                                # before the closing]]
                                tag = '[[FILE:' + identifier + '(' + comment + ')' + ']]'
                            else:
                                # No comment was present so just process the identifier

                                # Note: this logic is not sensitive to spaces around the
                                # Note by ravi : if the identifier is "file1 " then it can handle such cases.
                                # identifier
                                tag = '[[FILE:' + identifier + ']]'
                                text_add = substring.replace((identifier + ']]'), '')
                            if tag  in  run.text:
                                # Remove current tag from substring and then save remaining
                                # text of current substring
                                # print('idd',identifier)
                                try:
                                    # Retrieve the filepath for the current FILE identifier
                                    # and iterator value.  Exception raised if not found
                                    filepath = list(file_dataframe[identifier])[index]
                                    pass_replacements=pass_replacements+1
                                    # Retrieve a list of strings, reflecting unformatted
                                    # text from all Run instances for the Word file at the
                                    # provided filepath
                                    para_runs_for_filepath = filepaths_and_corresponding_lists_of_unformatted_strings[filepath]

                                    # ------------------------------------------------------
                                    # Insert text for identified file into spot of previous
                                    # [[FILE:...]] tag
                                    # ------------------------------------------------------
                                    for para_count,runs in para_runs_for_filepath.items():
                                        para=add_paragraph_after(para,runs)
                                    # # Iterate through all runs for file to insert
                                    # ------------------------------------------------------
                                    # Add content after the [[FILE...]] tag, in the original
                                    # document, after the content that replaced the [[FILE...]]
                                    # tag
                                    # ------------------------------------------------------
                                    sentence = new_paragraph.add_run(text_add)
                                    setAttributes(sentence, run)
                                except Exception as e:
                                    print(e)
                                    new_paragraph.add_run(tag)
                                    sentence = new_paragraph.add_run(text_add)
                                    setAttributes(sentence, run)
                                    print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))
                            else:
                                print('Error the Tag {} is not present in the run.'.format(tag))
                    else:
                        # We must have parsed all [[FILE...]] tags because
                        # no [[FILE...]] tags are currently present.  Therefore,
                        # create a Run instance from the unformatted text of the
                        # current Run instance and add the created Run instance
                        # to the paragraph
                        # Question: are we losing formatting here?
                        sentence= para.add_run(run.text)
                        setAttributes(sentence,run)
                # Remove current block from parent entity's container of
                # block instances
                delete_paragraph(block)
        # Process as python-docx.Table instance if it is one
        elif isinstance(block, Table):
            # print('Enter Table: ',table)
            # ll

            # Iterate through each python-docx.Row instance within
            # current python-docx.Table instance
            for row in block.rows:

                # Iterate through each python-docx.Cell instance within
                # current python-docx.Row instance
                for cell in row.cells:

                    # Iterate through each python-docx.Paragraph instance
                    # within current python-docx.Cell instance
                    for paragraph in cell.paragraphs:

                        # Extract unformatted text from current Paragraph instance
                        block_text = paragraph.text

                        # print('bt ',block_text)

                        # Retrieve Pandas data frame for FILES worksheet
                        # Action Item: convert 3 to constant defined up front
                        file_dataframe = dataframes[3]

                        # Check whether start of FILE tag is present within
                        # unformatted text for block.  Unformatted text is a
                        # viable way to check for the tag because we don't
                        # care about formatting considerations (distinguishing
                        # Runs) for this check
                        if '[[FILE:' in block_text:

                            # -------------------------------------------------------
                            # Add a paragraph after the current paragraph; taken from
                            # https://stackoverflow.com/questions/48663788/python-docx-insert-a-paragraph-after/53762537#53762537
                            # -------------------------------------------------------

                            # Create Paragraph element
                            new_p = OxmlElement("w:p")

                            # Add new Paragraph instance to current Paragraph instance
                            paragraph._p.addnext(new_p)

                            # Create new Paragraph instance for new_p with
                            # paragraph's parent as parent element
                            new_para = Paragraph(new_p, paragraph._parent)
                            para=new_para
                            # Iterate through each Run instance within a
                            # Paragraph instance
                            for run in paragraph.runs:

                                # Is [[FILE: within unformatted text of current Run instance
                                if '[[FILE:' in run.text:
                                    # Extract substrings before and after [[FILE:
                                    substrings = (run.text).split('[[FILE:')
                                    # Add a Run instance, containing the unformatted
                                    # text preceding the [[FILE: text, to the Paragraph instance
                                    sentence=new_para.add_run(substrings[0])
                                    setAttributes(sentence,run)

                                    # Iterate through all substrings (excluding the initial
                                    # substring because it precedes the first tag)
                                    for substring in substrings[1:]:

                                        # Extract the [[FILE...]] tag's identifier
                                        # (which indexes a filepath in the data frame)
                                        # from the current substring
                                        identifier = substring.split(']]')[0]

                                        # Handle comment if present
                                        if '(' in identifier:
                                            identifier = identifier.split('(')[0]
                                            comment = identifier.split('(')[1].split(')')[0]
                                            tag = '[[FILE:' + identifier + '(' + comment + ')' + ']]'
                                        # Action Item: does not handle spaces before or after
                                        # parentheses
                                        else:
                                            tag = '[[FILE:' + identifier + ']]'
                                        # Does not handle spaces before/after identifier
                                        if tag in run.text:

                                            # Determine the text to add as the current substring
                                            # after the remaining part of the tag is removed.
                                            # Action Item: is not resilient for spaces around identifier
                                            text_add = substring.replace((identifier + ']]'), '')
                                            # print('idd',identifier)

                                            # Recreate full FILE tag from identifier

                                            # print('filename',file_df[identifier])

                                            # Retrieve filepath for current identifier and current
                                            # index value from files data frame (FILES workbook
                                            # data)
                                            try:
                                                # Retrieve filepath for current identifier and index; raises
                                                # exception if not present
                                                filepath = list(file_dataframe[identifier])[index]
                                                pass_replacements=pass_replacements+1
                                                # Retrieve a list of strings, each reflecting the unformatted
                                                   # text content for all Runs in the identified filepath
                                                runs_for_filepath = \
                                                filepaths_and_corresponding_lists_of_unformatted_strings[filepath]

                                                for para, runs in para_runs_for_filepath.items():
                                                    add_paragraph_after(para, runs)
                                                    # sen.style=current_run.style
                                                # Now add text for the remainder after the last tag
                                                sentence = para.add_run(text_add)
                                                setAttributes(sentence, run)
                                            except Exception as e:
                                                # Handle error - no identifier present in Globals workbook
                                                new_para.add_run(tag)
                                                sentence = new_para.add_run(text_add)
                                                setAttributes(sentence, run)
                                                print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))
                                        else:
                                            print('Error the Tag {} is not present in the run.'.format(tag))
                                else:
                                    sentence = para.add_run(run.text)
                                    setAttributes(sentence, run)
                            # Set instance flag reflecting that [[FILE...]] tag
                            # replacement has occurred
                            # Recover Paragraph instance
                            delete_paragraph(paragraph)
                        # print('Replaced: ',paragraph.text)
        else:
            print("Warning: current block is neither a Paragraph nor a Table instance!")
    if pass_replacements>0:
        replacement_done=True
    # Return python-docx.Document instance with zero or one [[FILE...]] tag replaced
    # with content from corresponding Word file (if replacement occurred) and instance
    # flag as True if replacement occurred
    return document, replacement_done


def generate_filepaths_and_corresponding_lists_of_unformatted_strings(filepaths):  # CONFIRMED
    """
    Creates a dictionary of filepath-content_string pairs from
    keys for each filepath in the list of provided filepaths and
    values as a string comprising all unformatted content from
    Paragraphs only within documents at the provided filepaths
    :param filepaths: (list of string) filepaths for Word files
    :return: (dict of string-string pairs) filepath keys and
    unformatted string of Paragraph content from filepath values
    """
    # Create an empty dictionary for receiving file data
    filepaths_and_corresponding_lists_of_unformatted_strings = {}

    # Iterate through each filepath received as ARGV
    for filepath in filepaths:
        # Create a python-docx.Document instance for current
        # filepath and read corresponding Word file into instance
        document = Document(filepath)

        # Generate a list of unformatted text strings, for each run
        # in the document, and store that list as a value that is keyed
        # on the filepath for the document.  This scheme allows for an
        # easy retrieval of unformatted run text for a document from the
        # document's filepath
        filepaths_and_corresponding_lists_of_unformatted_strings[filepath] = \
            identify_unformatted_runs_from_document(document)

    # print(filepaths_and_corresponding_lists_of_unformatted_strings)
    return filepaths_and_corresponding_lists_of_unformatted_strings


def preprocess_files(input_word_filepath, dataframes,iterator):  # CONFIRMED3
    """
    Handles [[IF...]] and [[FILE...]] tag processing within
    input_word_filepath, so that [[IMAGE...]] and [[TEXT...]]
    content can be inserted in the process_files() function.
    :param input_word_filepath: (string) filepath for input Word file
    that will be processed
    :param dataframes: (list of Pandas data frames) workbook data
    :return: (empty)
    """

    # --------------------------------------------------------------
    # Read Word files at filepaths included within [[FILE...]] tags
    # --------------------------------------------------------------

    # Generate a list of all unique filepaths for all FILE tags
    # for all iterator values.  This list will be used for parsing
    # all files upfront, and only once, notwithstanding that those
    # files could be included in several ways and in several
    # locations in other files (efficient Word file parsing)
    #
    # A key aspect of this functionality working is that identifiers
    # in [[IF...]] tags may evaluate as TRUE or FALSE for different
    # iterators so correct [[IF...]] tag replacement depends on
    # evaluating the [[IF...]] tags on a per-iterator basis not on
    # a global basis
    unique_filepaths_for_file_tags = identify_unique_filepaths_from_all_file_tags(dataframes)

    # Check that input Word file is not in list of filepaths because
    # that situation would reflect a circular FILE tag situation
    if input_word_filepath in unique_filepaths_for_file_tags:
        print("ERROR: input Word file is also a filepath for a FILE tag - circular FILE problem")
        exit(1)

    # Generate a dictionary, from all unique filepaths for
    # all file tags, of each file path as a key and values as a
    # list of strings, reflecting all unformatted content for
    # all Run instances within the document
    filepaths_and_corresponding_lists_of_unformatted_strings = generate_filepaths_and_corresponding_lists_of_unformatted_strings(unique_filepaths_for_file_tags)

    # Add input Word file to end of list of unique filepaths
    unique_filepaths_for_file_tags.append(input_word_filepath)

    # Iterate through each filepath for all FILE tags across all
    # iterators (plus input Word file)
    for index in range(len(unique_filepaths_for_file_tags)):

        # Retrieve current filepath from list of unique filepaths
        filepath = unique_filepaths_for_file_tags[index]

        # Determine if current file is last file in list (would
        # be the input Word file) because input Word file was appended
        # to list
        if index == len(unique_filepaths_for_file_tags) - 1:
            # Because the input Wordfile was originally omitted
            # from the list of filepaths when lists of unformatted
            # content was retrieved for each filepath, we must
            # generate the list of unformatted content for the
            # input Word file.
            #
            # This was done by regenerating lists of unformatted content
            # strings for all filepaths (i.e., filepaths for which the
            # lists of content strings had been generated plus the list
            # for the input Wordfile).  Most of that regeneration was
            # redundant.  Therefore, the lines with redundancy have been
            # commented out and lines were added that add a list of
            # strings for the input Word file to the collection
            #
            # An open question is whether we can simplify this further by
            # adding the input Word file to the list of filepaths prior
            # to generating any lists of content strings.  In the spirit
            # of avoiding changes to logic, however, that change will not
            # be proposed now

            # filepaths_and_corresponding_lists_of_unformatted_strings = \
            #	generate_filepaths_and_corresponding_lists_of_unformatted_strings(unique_filepaths_for_file_tags)

            # generate a list of unformatted content strings for input Wordfile
            additional_content_strings_for_input_wordfile = generate_filepaths_and_corresponding_lists_of_unformatted_strings([input_word_filepath])

            # Add the list of unformatted strings, for the input Word file,
            # to the dictionary of all filepath-stringlist pairs
            filepaths_and_corresponding_lists_of_unformatted_strings[input_word_filepath] = additional_content_strings_for_input_wordfile[input_word_filepath]

        # --------------------------------------------------------------
        # Read input Word file
        # --------------------------------------------------------------

        # Read the Word document at the current filepath
        document = Document(unique_filepaths_for_file_tags[index])

        # --------------------------------------------------------------
        # Replace [[IF...]] tags in input Word file
        # --------------------------------------------------------------

        print("WARNING: skipping IF tag parsing and replacement")
        #  New logic for IF ELSE nested tags replacement which strictly assumes that IF or ELSE tag should completely fall inside a single run as per the original specs
        #  of the document.
        # replace_nested_tags(dataframes,filepath,iterator)
        # replace [[IF...]] tags as appropropriate for iterator
        # GLOBAL values and store result at filepath
        # find_and_replace_if_tags(dataframes,
        #                          document,
        #                          index,
        #                          unique_filepaths_for_file_tags)

        # --------------------------------------------------------------
        # Reread input Word file after replacing [[IF...]] tags
        # --------------------------------------------------------------

        # Reread Word file at filepath into a python-docx.Document
        document = Document(unique_filepaths_for_file_tags[index])
        # content=identify_unformatted_runs_from_document(document)

        # --------------------------------------------------------------
        # Replace all [[FILE...]] tags within input Word file
        # --------------------------------------------------------------

        # Repeatedly call replace_file_tags() until that function returns false,
        # indicating that all [[FILE...]] tags have been replaced
        while True:
            # Replace [[FILE...]] tags within document with content from
            # corresponding Word files
            document, tag_replacement_performed = replace_file_tags(dataframes,
                                                                    document,
                                                                    filepaths_and_corresponding_lists_of_unformatted_strings,
                                                                    iterator)

            # print('contain: ',tag_replacement_performed)

            # Determine whether looping should continue based on flag
            # returned from replace_file_tags()
            if tag_replacement_performed == True:
                pass
            else:
                break

        #  This code your wrote is not correct  since it returns the first instance of the document. what the idea here is to save each file we open we save it so that the next time we read it
        #  we hve the updted one. If we return from here then the function won't get completed.
        # # [[FILE...]] tag replacement is complete so return resulting Document instance
        # return document

        # --------------------------------------------------------------
        # Store modified Word file
        # --------------------------------------------------------------

        document.save(filepath)


# return document


def validatesheetnames(sheet_names):  # CONFIRMED
    """
    Validates whether list of sheet names is valid (includes all required
    worksheet names and no extra workheet names)
    :param sheet_names: (list of strings) worksheet names
    :return: (boolean): True if the sheet_names list reflects a correct
    four sheet names; False otherwise
    """

    if sheet_names[0] == 'GLOBALS' and sheet_names[1] == 'TEXT' and sheet_names[2] == 'IMAGE' and sheet_names[3] == 'FILE':
        return True
    return False


def replacehead(df):# Confirmed
    """
    Extracts the column headers from the first row of the dataframe
    and adjusts dataframe to remove column header row but set column
    headings from that row's contents
    :param df: (Pandas data frame) Pandas representation of worksheet
    :return: (Pandas data frame) modified data frame
    """

    # Return the header row from the worksheet as a series of key-value pairs
    new_header = df.iloc[0]  # grab the first row for the header

    # Remove the header row from the dataframe
    df = df[1:]  # take the data less the header row

    # Sets the column labels for the dataframe to be the
    # values from the extracted header row
    df.columns = new_header

    # Return modified dataframe
    return df


def cleanExcel(dataframe):  # CONFIRMED
    """
    Extracts dataframes for individual worksheets, removes header rows from
    those dataframes and populates column header labels from that row's data,
    and removes empty cells
    :param dataframe: (Pandas data frame) workbook data (multiple worksheets)
    :return: (list of four Pandas data frames) data frames with header rows
    removed but column headers configured from those rows and empty cells removed
    """

    # Remove the header row from the dataframe and configure column headings
    # to reflect the contents of the header row through the replacehead function.
    # Additionally, the "dropna" function removes empty cells; see
    # https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.dropna.html#pandas.DataFrame.dropna
    dataframe_globals = replacehead(dataframe.parse('GLOBALS').dropna(how='all').dropna(axis='columns'))
    dataframe_text = replacehead(dataframe.parse('TEXT').dropna(how='all').dropna(axis='columns'))
    dataframe_image = replacehead(dataframe.parse('IMAGE').dropna(how='all').dropna(axis='columns'))
    dataframe_file = replacehead(dataframe.parse('FILE').dropna(how='all').dropna(axis='columns'))
    # print(dataframe_file)

    # Return list of 4 dataframes as processed above
    return dataframe_globals, dataframe_text, dataframe_image, dataframe_file


def read_excel(filepath):  # CONFIRMED
    """
    Read Excel workbook, whose filepath is specified as file_name, into Pandas
    and return a list of Pandas dataframes - one each for globals, text tags,
    image tags and file tags or return 0 if Excel workbook not parsed properly.
    Removes header rows, stores header labels in dataframes, and removes empty
    cells
    :param file_name: (string) filepath for Excel workbook
    :return: (list of dataframes or 0 if error) list of dataframes for globals,
    text tags, image tags and file tags after removing header row, setting
    header labels, and removing empty cells
    """

    # Read specified Excel workbook into dataframe;
    # see https://pandas.pydata.org/pandas-docs/stable/reference/io.html#excel
    dataframe = pandas.ExcelFile(filepath)

    # print(dataframe.sheet_names)

    # Confirm that worksheets read from Excel file are complete and correct;
    # returns True if they are complete and correct
    if validatesheetnames(dataframe.sheet_names):
        # Worksheet names were complete and correct

        # Extract the individual worksheets from the dataframe
        # read from the Excel workbook, remove header row and store
        # header labels in dataframe, and remove empty cells
        dataframe_globals, dataframe_text, dataframe_image, dataframe_file = cleanExcel(dataframe)
        # Return list of 4 dataframes as modified above
        return [dataframe_globals, dataframe_text, dataframe_image, dataframe_file]
    else:
        # Worksheet names were not complete and correct to exit reflecting error
        return 0

def replace_text_tags(list_df, document, index, target_file):
    """
    Replaces [[TEXT:<identifier>]] tags with text corresponding to the identifier
    as specified in the TEXT worksheet and for the current index.  Uses the
    "text" attribute of a Paragraph instance to retrieve the unformatted text
    for the paragraph.  We are able to use the unformatted text version of the
    paragraph because we require that formatting is consistent for a tag, which
    results in tags that do not span a "Run" in python-docx parlance
    :param list_df: list of Pandas dataframes reflecting worksheets
    for globals, text tags, image tags and file tags
    :param document: (python-docx.Document) Document instance
    :param index: (integer) current index value
    :return: (python-docx.Document) modified Document instance
    """

    # Iterate through each Paragraph and Table within Document
    # based on iter_block_items(), which emits Paragraph and
    # Table instances within the Document
    for block in iter_block_items(document):

        # Process block as paragraph if current block is an instance of Paragraph
        if isinstance(block, Paragraph):

            # Extract unformatted text from Paragraph instance
            block_text = block.text
            # print(block_text)

            # Determine whether [[TEXT:<identifier>(optional-comment)]] is present
            if '[[TEXT:' in block_text:

                # At least one [[TEXT...]] tag is present so find substrings for
                # the one or more [[TEXT...]] tags in the current paragraph.  Returns
                # a list of substrings separated on '[[TEXT:'' such that the initial
                # element is the string preceding the first [[TEXT...]] tag, the
                # next element is text after [[TEXT: until the start of the next
                # [[TEXT...]] tag or the end of line, repeating thereafter.  Thus,
                # the first characters in the first element are the identifier and
                # optional comment for the [[TEXT...]] tag
                new_p = OxmlElement("w:p")
                block._p.addnext(new_p)
                new_para = Paragraph(new_p, block._parent)
                for run in block.runs:
                    if '[[TEXT:' in run.text:
                        numtexttags = (run.text).split('[[TEXT:')
                        new_para.add_run(numtexttags[0])
                        # Iterate through each [[TEXT...]] tag, after skipping text preceding
                        # the first [[TEXT...]] tag (the 1: expression)
                        for tag in numtexttags[1:]:
                            # for current substring, split the end of the tag from text
                            # after the end of the tag.  The identifier is to the left of
                            # the initial substring
                            identifier = tag.split(']]')[0]
                            if '(' in identifier:
                                identifier = identifier.split('(')[0]
                                comment = identifier.split('(')[1].split(')')[0]
                                tag_d = '[[TEXT:' + identifier + '(' + comment + ')' + ']]'
                            else:
                                tag_d = '[[TEXT:' + identifier + ']]'
                            text_add = tag.replace((identifier + ']]'), '')
                            # Recreate tag so it can be replaced easily in Paragraph instance
                            # tag_d='[[TEXT:'+identifier+']]' # Question2: will this later include comment parsing?
                            #  try catch Required to avoid key error in pandas and is used to validate and detect invalid tag.If you want to change it we can discuss
                            if tag_d in run.text:
                                try:
                                    # print('Tag ID : ',tag_d)
                                    text_rep = list(list_df[1][identifier])[index]
                                    new_para.add_run(text_rep)
                                    sentence=new_para.add_run(text_add)
                                    setAttributes(sentence,run)
                                # Replace [[TEXT...]] tag with its corresponding text
                                # from TEXT worksheet
                                # block_text=block_text.replace(tag_d,text_rep)
                                # print('block_text: ',block_text)
                                except Exception as e:
                                    new_para.add_run(tag_d)
                                    sentence=new_para.add_run(text_add)
                                    setAttributes(sentence,run)
                                    print('Error the Tag {} is not present in the Excel Sheet.'.format(tag_d))
                            else:
                                print('Error the Tag {} is not present in the Run.'.format(tag_d))
                    else:
                        sentence = new_para.add_run(run.text)
                        setAttributes(sentence, run)
                delete_paragraph(block)
            else:
                # No [[TEXT...]] tag present in paragraph
                pass
        # Process block as table if current block is an instance of Table
        elif isinstance(block, Table):
            # Iterate through each Row instance within the Table
            for row in block.rows:
                # Iterate through each Cell instance within the Row
                for cell in row.cells:
                    # Iterate through each Paragraph instance within the Cell
                    for paragraph in cell.paragraphs:
                        # Extract unformatted text from Paragraph instance
                        block_text = paragraph.text

                        # NOTE THAT THIS CODE REPEATS FUNCTIONALITY FOR PARAGRAPHS

                        # Future project: break logic below into function that is
                        # called for Paragraph instances and Table instances

                        # Detect presence of [[TEXT...]] tag in current block
                        if '[[TEXT:' in block_text:
                            # [[TEXT...]] tag found so parse it

                            # Create a list of substrings, the first of which
                            # ends just before the first [[TEXT:...]] tag, and
                            # the remainder of which begins at the start of each
                            # [[TEXT...]] tag
                            new_p = OxmlElement("w:p")
                            paragraph._p.addnext(new_p)
                            new_para = Paragraph(new_p, paragraph._parent)
                            for run in paragraph.runs:
                                if '[[TEXT:' in run.text:
                                    numtexttags = (run.text).split('[[TEXT:')
                                    new_para.add_run(numtexttags[0])
                                    # Iterate through each [[TEXT...]] tag, after skipping text preceding
                                    # the first [[TEXT...]] tag (the 1: expression)
                                    for tag in numtexttags[1:]:
                                        # for current substring, split the end of the tag from text
                                        # after the end of the tag.  The identifier is to the left of
                                        # the initial substring
                                        identifier = tag.split(']]')[0]
                                        if '(' in identifier:
                                            identifier = identifier.split('(')[0]
                                            comment = identifier.split('(')[1].split(')')[0]
                                            tag_d = '[[TEXT:' + identifier + '(' + comment + ')' + ']]'
                                        else:
                                            tag_d = '[[TEXT:' + identifier + ']]'
                                        if tag_d in run.text:
                                            text_add = tag.replace((identifier + ']]'), '')
                                            # Recreate tag so it can be replaced easily in Paragraph instance
                                            # tag_d='[[TEXT:'+identifier+']]' # Question2: will this later include comment parsing?
                                            #  try catch Required to avoid key error in pandas and is used to validate and detect invalid tag.If you want to change it we can discuss
                                            try:
                                                # print('Tag ID : ',tag_d)
                                                text_rep = list(list_df[1][identifier])[index]
                                                new_para.add_run(text_rep)
                                                sentence = new_para.add_run(text_add)
                                                setAttributes(sentence, run)
                                            # Replace [[TEXT...]] tag with its corresponding text
                                            # from TEXT worksheet
                                            # block_text=block_text.replace(tag_d,text_rep)
                                            # print('block_text: ',block_text)
                                            except Exception as e:
                                                new_para.add_run(tag_d)
                                                print('Error the Tag {} is not present in the Excel Sheet.'.format(tag_d))
                                        else:
                                            print('Error the Tag {} is not present in the Run.'.format(tag_d))
                                else:
                                    sentence = new_para.add_run(run.text)
                                    setAttributes(sentence, run)
                                    print('Error the Tag {} is not present in the Run.'.format(tag_d))
                            delete_paragraph(paragraph)
        else:
            print("Warning: Current block is neither a Paragraph nor a Table")

    # Return Document instance to caller so next processings step can build on
    # this result
    document.save(target_file)
    return document

def insert_image_after(run, image, img_width=-1, img_height=-1):
    """
    Inserts the provided image into the provided run, optionally with
    width and height dimensions, and directs python-docx to scale based
    on width and/or height dimensions
    :param run: (python-docx.Run) Run instance into which image is inserted
    :param image: (string) filepath of image to insert
    :param img_width: (float) width of image in inches
    :param img_height: (float) height of image in inches
    :return: (empty)
    """

    if img_height == -1:
        # image height not specified so check image width
        if img_width == -1:
            # image height and width not specified so let python-docx
            # determine image dimensions
            run.add_picture(image)
        else:
            # image width but not height specified so tell python-docx
            # to scale the image proportionally for the width
            run.add_picture(image, Inches(img_width))
    else:
        # image height specified so check image width
        if img_width == -1:
            # image height specified but not image width so tell
            # python-docx to scale image proportionally for the height
            run.add_picture(image, Inches(img_height))
        else:
            # image width and height provided so tell python-docx
            # to scale the image as provided in both dimensions
            run.add_picture(image, width=Inches(img_width), height=Inches(img_height))


def insert_run_after(paragraph, text, style=None):
    """
    Uses internal python-docx methods to add a new Run instance, containing the
    provided text, at the end of the provided Paragraph instance

    See https://stackoverflow.com/questions/48663788/python-docx-insert-a-paragraph-after
    :param paragraph: (python-docx.Paragraph) a paragraph
    :param text: (string) text to add
    :param style: (not currently implemented) see https://python-docx.readthedocs.io/en/latest/api/style.html
    :return: (empty)
    """

    # Create a new data structure for receiving text and adding data
    # structure to received paragraph
    new_p = OxmlElement("w:p")

    # Add new data structure to received paragraph
    paragraph._p.addnext(new_p)

    # Create new Paragraph instance linked to parent Paragraph instance
    # using new data structure
    new_para = Paragraph(new_p, paragraph._parent, style)

    # Add text to new Paragraph instance as a new Run instance
    new_para.add_run(text)


def delete_paragraph(paragraph):  # CONFIRMED
    """
    Remove received Paragraph instance from its patent's
    collection of Paragraph instances, thereby removing the
    Paragraph instance from its contained document
    :param paragraph: (python-docx.Paragraph) paragraph
    :return: (empty)
    """

    # Retrieve content container for Paragraph instance
    p = paragraph._element

    # Remove current Paragraph instance from Paragraph instances
    # stored in parent element
    p.getparent().remove(p)

    # Set paragraph content to None so memory is recovered
    p._p = p._element = None


def replace_image_tag(list_df, document, index, target_file):
    """
    Replaces [[IMAGE:<identifier>]] tags with an image corresponding to the identifier
    as specified in the IMAGES worksheet and for the current index.  Uses the
    "text" attribute of a Run instance to retrieve the unformatted text
    for the paragraph.  We are able to use the unformatted text version of the
    paragraph because we require that formatting is consistent for a tag, which
    results in tags that do not span a "Run" in python-docx parlance
    :param list_df: list of Pandas dataframes reflecting worksheets
    for globals, text tags, image tags and file tags
    :param document: (python-docx.Document) Document instance
    :param index: (integer) current index value
    :return: (python-docx.Document) modified Document instance
    """

    # Iterate through each Paragraph and Table within Document
    # based on iter_block_items(), which emits Paragraph and
    # Table instances within the Document
    for block in iter_block_items(document):

        # Process block as paragraph if current block is an instance of Paragraph
        if isinstance(block, Paragraph):
            # Extract unformatted text from Paragraph instance
            block_text = block.text

            if '[[IMAGE:' in block.text:
                # Iterate through each Run within current block
                # (a Paragraph instance)
                new_p = OxmlElement("w:p")
                block._p.addnext(new_p)
                new_para = Paragraph(new_p, block._parent)
                for run in block.runs:
                    # print('run : ',run.text)
                    try:
                        # Determine whether [[IMAGE:<identifier>(optional-comment)]] is present
                        if '[[IMAGE:' in run.text:
                            # At least one [[IMAGE...]] tag is present so find substrings for
                            # the one or more [[IMAGE...]] tags in the current run.  Returns
                            # a list of substrings separated on '[[IMAGE:'' such that the 0th
                            # element is the string preceding the first [[IMAGE...]] tag, the
                            # 1st element is text after [[IMAGE: until the start of the next
                            # [[IMAGE...]] tag or the end of line, repeating thereafter.  Thus,
                            # the first characters in the first element are the identifier and
                            # optional comment for the [[IMAGE...]] tag

                            image_items = (run.text).split('[[IMAGE:')
                            # print('image_items: ',image_items)
                            sentence= new_para.add_run(image_items[0])
                            setAttributes(sentence,run)
                            for image in image_items[1:]:
                                # print('split image : ',image.split(']]'))
                                # image_tag = '[[IMAGE:' + image.split(']]')[0] + ']]'
                                identifier = image.split(']]')[0]
                                if '(' in identifier:
                                    identifier = identifier.split('(')[0]
                                    comment = identifier.split('(')[1].split(')')[0]
                                    image_tag = '[[IMAGE:' + identifier + '(' + comment + ')' + ']]'
                                else:
                                    image_tag = '[[IMAGE:' + identifier + ']]'
                                text_add = image.replace((identifier + ']]'), ' ')
                                if image_tag in run.text:
                                    try:
                                        img_path = list(list_df[2][identifier])[index]
                                        img_width = list(list_df[2][identifier + '_width'])[index]
                                        img_height = list(list_df[2][identifier + '_height'])[index]
                                        insert_image_after(new_para.add_run(), img_path, img_width, img_height)
                                        sentence=new_para.add_run(text_add)
                                        setAttributes(sentence,run)
                                    except Exception as e:
                                        print(e)
                                        new_para.add_run(image_tag)
                                        sentence=new_para.add_run(text_add)
                                        setAttributes(sentence,run)
                                        print('Error the Tag {} is not present in the Excel Sheet.'.format(image_tag))
                                else:
                                    print('Error the Tag {} is not present in the Excel Sheet.'.format(image_tag))
                        else:
                            sentence = new_para.add_run(run.text)
                            setAttributes(sentence, run)
                    except Exception as e:
                        # print('ErrorPar: ',str(e))
                        print('Image Tag does not fall under  Run.')
                    # run.text=''
                delete_paragraph(block)
        # Process block as table if current block is an instance of Table
        elif isinstance(block, Table):
            # Iterate through each Row instance within the Table
            for row in block.rows:
                # Iterate through each Cell instance within the Row
                for cell in row.cells:
                    # Iterate through each Paragraph instance within the Cell
                    for paragraph in cell.paragraphs:
                        # Extract unformatted text from Paragraph instance
                        block_text = paragraph.text

                        if '[[IMAGE:' in paragraph.text:
                            new_p = OxmlElement("w:p")
                            paragraph._p.addnext(new_p)
                            new_para = Paragraph(new_p, paragraph._parent)
                            for run in paragraph.runs:
                                # print('run: ',run.text)
                                try:
                                    if '[[IMAGE:' in run.text:
                                        image_items = (run.text).split('[[IMAGE:')
                                        # print('image_items: ',image_items)
                                        sentence = new_para.add_run(image_items[0])
                                        setAttributes(sentence, run)
                                        for image in image_items[1:]:
                                            # print('split image : ',image.split(']]'))
                                            image_tag = '[[IMAGE:' + image.split(']]')[0] + ']]'
                                            identifier = image.split(']]')[0]
                                            if '(' in identifier:
                                                identifier = identifier.split('(')[0]
                                                comment = identifier.split('(')[1].split(')')[0]
                                                image_tag = '[[IMAGE:' + identifier + '(' + comment + ')' + ']]'
                                            else:
                                                image_tag = '[[IMAGE:' + identifier + ']]'

                                            text_add = image.replace((identifier + ']]'), ' ')
                                            # print('text_add: ' ,text_add)
                                            if image_tag in run.text:
                                                try:
                                                    img_path = list(list_df[2][identifier])[index]
                                                    img_width = list(list_df[2][identifier + '_width'])[index]
                                                    img_height = list(list_df[2][identifier + '_height'])[index]
                                                    insert_image_after(new_para.add_run(), img_path, img_width, img_height)
                                                    sentence = new_para.add_run(text_add)
                                                    setAttributes(sentence, run)
                                                except Exception as e:
                                                    print(e)
                                                    new_para.add_run(image_tag)
                                                    sentence = new_para.add_run(text_add)
                                                    setAttributes(sentence, run)
                                                    print('Error the Tag {} is not present in the Excel Sheet.'.format(
                                                        image_tag))
                                            else:
                                                print('Error the Tag {} is not present in the Run.'.format(
                                                    image_tag))
                                    else:
                                        sentence = new_para.add_run(run.text)
                                        setAttributes(sentence, run)
                                except Exception as e:
                                    print('Image Tag Doesnt fall under run ')
                            delete_paragraph(paragraph)
    document.save(target_file)
    return document  # with modifications


def process_document(dataframes, preprocessed_document, index, target_file, input_file):
    """
    Receives a Document instance that has been preprocessed (all [[FILE...]]
    and [[IF...]] tags have been replaced) and processes that Document
    instance by replacing [[IMAGE...]] and [[TEXT...]] tags with their
    corresponding content
    :param dataframes: (list of Pandas dataframes) comprises data frames
    for globals, text tags, image tags and file tags from spreadsheet
    :param preprocessed_document: (python-docx.Document) Document
    instance storing contents of input Word file that will be modified
    but which also has [[FILE...]] and [[IF...]] tags already processed
    :param index: (int) current index value (determines which row of
    Excel data will be used)
    :return: (python-docx.Document) Document instance read from
    input Word file after parsing for all IF/FILE/IMAGE/TEXT tags
    """

    # Parse Document instance for IF-ELSE-ENDIF tags and replace those
    # tags with appropriate conditional content

    print("Warning: IF tag processing is commented out")
    #  New logic for IF ELSE nested tags replacement which strictly assumes that IF or ELSE tag should completely fall inside a single run as per the original specs
    #  of the document.
    # replace_nested_tags(dataframes,filepath,iterator)

    preprocessed_document = Document(input_file)

    # Parse Document instance for TEXT tags and replace those
    # tags with appropriate text
    replace_text_tags(dataframes, preprocessed_document, index, target_file)

    preprocessed_document = Document(target_file)

    # Parse Document instance for IMAGE tags and replace those
    # tags with appropriate image content
    replace_image_tag(dataframes, preprocessed_document, index, target_file)


# return document_after_image_tag_replacement
# return document_after_text_tag_replacement


if __name__ == '__main__':

    # ----------------------------------------------------------------
    # Processing Excel Workbook
    # ----------------------------------------------------------------

    print('Loading And Validating EXCEL')

    # Create filepath for Excel workbook as CWD/Files/<workbook>
    excel_filepath = os.path.join(os.path.join(os.getcwd(), 'Files'), 'excelfile.xlsx')

    # Return a list of Pandas dataframes for globals, text tags, image tags and file tags
    # from separate worksheets (for each) within Excel workbook after removing
    # header row, storing header labels in data frames, and removing empty cells
    dataframes = read_excel(excel_filepath)

    # Handle error from Pandas processing of Excel workbook (a 0 returned
    # reflects an error in dataframe parsing (workbook parsing))
    if type(dataframes) is int:
        print("Error reading Excel workbook at %s.  Exiting." % excel_filepath)
        exit(1)

    print("Excel validation complete.")

    # ----------------------------------------------------------------
    # Identifying input Word filepath and destination Word filepaths
    # ----------------------------------------------------------------

    # Create filepath for input document as CWD/Files/<input-file>
    input_word_filepath = os.path.join(os.path.join(os.getcwd(), 'Files'), 'inputWordfile.docx')

    ## NOT USED
    # generate_tags_list(input_word_filepath,dataframes)

    # Generate a list of Word files to be generated - from the GLOBALS
    # worksheet because the dataframe for the GLOBALS worksheet is at
    # index=0 and from the DESTINATION column within that worksheet
    target_word_filepaths = list(dataframes[0]['DESTINATION'])
    # print(target_file)
    # ll

    # ----------------------------------------------------------------
    # Process the input Word file to produce each of the target Word
    # files by conditionally including content via [[IF:...]] tags and
    # by including FILE/IMAGE/TEXT content dependent in part on iterator
    # for each target Word file
    # ----------------------------------------------------------------

    # Iterate through each Word file to be generated (each target Word filepath)
    for index in range(len(target_word_filepaths)):
        ## Retrieve name of current Word file to be created
        target_word_filepath = target_word_filepaths[index]
        print('Preprocessing file %s to generate %s' % (input_word_filepath, target_word_filepath))

        # Processes [[IF...]] and [[FILE...]] tags to produce Document
        # instances reflecting the result of that processing (although
        # [[IF...]] tag processing is a work in progress so I commented
        # it out
        #
        # Reads, preprocesses and writes Word file to input_word_filepath
        preprocess_files(input_word_filepath, dataframes,index)

        print('Iteration ', index + 1)

        # Read preprocessed document, generated by preprocess_files(),
        # from input_wordfile
        preprocessed_document = Document(input_word_filepath)

        # Find and operate on [[IMAGE...]] and [[TEXT...] tags to
        # produce a modified Document instance
        processed_document = process_document(dataframes,
                                              preprocessed_document,
                                              index,
                                              target_word_filepath,
                                              input_word_filepath)

    # Now save the result as the target Word file
    # processed_document.save(target_Word_filepath)

# read_document(list_df,input_wordfile,target_word_filepath[0],0)