from docx import Document
import pandas as pd
import os 
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
import time
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches

# Dave's most important comments
#
# I understand that [[IF...]] tag processing is still in progress so I did
# not look over that code
#
# Tag replacement - including [[FILE...]] tag replacement - seems to lose
# all formatting information (except for [[IMAGE...]] tag replacement as
# discussed below) because tag replacement relies on extracting unformatted
# content for a Paragraph, modifying that content, and storing the modified
# content into the Paragraph instance.  Thus, any formatting within the
# Paragraph will be lost.  This general approach should work, without
# losing formatting, if this operation is done on a Run level rather than
# a Paragraph level, because formatting is unchanged within a Run.
# However, [[IMAGE...]] tag replacement operates at a Run level and seems
# to avoid this issue, thereby suggesting that if Ravi confirms the issue
# with formatting loss, a solution is already shown for the issue
#
# Tag parsing does not consider the presence of optional comments surrounded
# by parentheses; e.g., [[FILE=identifer99 (this is the main file)]].  This
# feature is not high priority and can be deferred for now (possibly to a
# follow-on project that would expand the current project)
#
# I understand that Excel file processing and validation is in progress
#
# Note that I commented out logic for [[IF...]] tag processing and did
# not review that code because I understand that nested [[IF...]] tag
# processing is in progress
#
# The functions below are not tied into the executing code at present, so I
# have not reviewed them (with the caveat that I chose not to review code
# for [[IF...]] tag processing because I understand that logic for handling
# nested [[IF...]] tags is in progress.
# - generatefiletags
# - generate_tags_list
# - evaluateTag
# - replaceIftext
# - find_and_replace_if_tags
# - validatesheetnames
# - replacehead
# - cleanExcel
# - read_excel
# - validate_excel



def iter_block_items(parent):
	"""
	Provides an iterator interface for retrieving python-docx Paragraph and
	Table entities from a parent/containing entity
	:param parent: (python.docx instance of some type) parent entity
	:return: (iterator output) emits a series of python-docx.Parent or Table
	instances based on the content of the parent entity
	"""

	# Performs type-specific processing, based on data type of parent, to
	# extract parent-data-type-specific data from the parent such as paragraphs
	# or tables from a Document.  This step sets up an easy capability to iterate
	# through the children of parent on a parent-type-independent basis
	if isinstance(parent, _Document):
		# parent is a python-docx Document so set parent element to the
		# document's body container
		parent_element = parent.element.body
	elif isinstance(parent, _Cell):
		# parent is a python-docx table cell so set parent element to the container
		# for the table cell
		parent_element = parent._tc
	elif isinstance(parent, _Row):
		# parent is a python-docx table row so set parent element to the container
		# for the table row
		parent_element = parent._tr
	else:
		## parent is none of document, table cell or table row so error
		raise ValueError("something's not right")

	# Iterate through each child element within the parent element's container
	for child in parent_element.iterchildren():
		if isinstance(child, CT_P):
			# current child is a paragraph so emit a python-docx.Paragraph instance
			yield Paragraph(child, parent)
		elif isinstance(child, CT_Tbl):
			# current child is a table so emit a python-docx.Table instance
			yield Table(child, parent)


def generatefiletags(input_wordfile,tags,tags_dict):
	document=Document(input_wordfile)
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			block_text=block.text
			# print('bt ',block_text)
			for i in range(len(tags)):
				tag=tags[i]
				if tag in block_text:
					complete_tag=tag+(block_text.split(tag)[1].split(']]')[0])+']]'
					# print('cTag: ',complete_tag)
					if i in tags_dict:
						tags_dict[i].append(complete_tag)
					else:
						tags_dict[i]=[complete_tag]
		elif isinstance(block, Table):
			for row in block.rows:
				row_data = []
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						block_text=paragraph.text
						for i in range(len(tags)):
							tag=tags[i]
							if tag in block_text:
								complete_tag=tag+(block_text.split(tag)[1].split(']]')[0])+']]'
								if i in tags_dict:
									tags_dict[i].append(complete_tag)
								else:
									tags_dict[i]=[complete_tag]						
						row_data.append(paragraph.text)
				# print("\t".join(row_data))
	# print(tags_dict)
	return tags_dict



def generate_tags_list(input_wordfile,list_df):
	tags=['[[FILE:','[[IMAGE:','[[TEXT:','[[IF:']
	tags_dict={}
	tags_dict=generatefiletags(input_wordfile,tags,tags_dict)
	#  complete verification later.




def getallfiles(list_df):
	"""
	Determine a list of all filepaths, for all FILE tags, for all
	iterator values in the FILE worksheet.  This function reflects
	an optimization to parse Word files at these filepaths once
	up front rather than parsing those files potentially many times
	:param list_df: (list of strings) (list of strings) list of
	filepaths for destination files
	:return: (list of strings) a unique list of all filepaths for
	all FILE tags and iterator values
	"""

	# Extract FILES (the 3rd indexed) data frame from data frames
	# read from workbook
	file_df=list_df[3]

	# Retrieve columns from FILES worksheet (not including 0th
	# column, which is the iterator column)
	filecolumns=list(file_df.columns)[1:]

	# print(filecolumns)

	# Create an empty list for storing FILE filepaths prior to
	# appending individual results to that list
	total_files=[]

	# Iterate through each column in the FILE workbook so that
	# we can develop a list of all filepaths for all iterator
	# values
	for col in filecolumns:
		total_files=total_files+list(file_df[col])

	# print(total_files)

	# Remove duplicates in the list by converting the
	# list to a dictionary, with keys matching values for
	# each entry, and then converting the resulting
	# dictionary back to a list
	unique_total_files=list(dict.fromkeys(total_files))

	# print(total_files)
	return unique_total_files

def evaluateTag(list_df,tag,i):

	glob_df=list_df[0]
	tag_data = (tag.split(':')[1]).split(']]')[0]
	if '(' in tag_data:
		tag_data = tag_data.split('(')[0].strip()
	tag_value=list(glob_df[tag_data])[i]

	if tag_value==True:
		tag_value='TRUE'
	else:
		tag_value='FALSE'

	return tag_value


def replaceIftext(list_df,block_text,i):
	"""

	:param list_df: (list of Pandas dataframes) Workbook data for globals,
	text tags, image tags and file tags
	:param block_text: (string) unformatted text retrieved from current python-docx element
	:param i: (integer) current iterator value
	:return:
	"""
	replace_text=""
	if_count=0
	else_count=0
	endif_count=0
	waitforendif=0
	waitforelse=0
	leave_tags_detected=0
	if_tags=[]
	if_tags_order=[]
	else_tags_order=[]
	text_replace=[]
	other_tags=['[[ELSE]]','[[ENDIF]]']
	leave_tags=['[[TEXT:','[[IMAGE:']
	split_data=block_text.split('[[IF:')
	# print(split_data)
	for x in split_data[1:]:
		if_tags.append('[[IF:'+x.split(']]')[0]+']]')
	# print(if_tags)
	dtect_tag=''
	leave_tag_text=''
	start=0
	tag_text=''
	z='TRUE'
	# print(block_text)
	for x in block_text:
		# print(x)
		if x=='[':
			start=1
		if start==1:
			dtect_tag=dtect_tag+x
			
		if start==0:
			tag_text=tag_text+x

		if dtect_tag in if_tags:
			if waitforendif==1:
				start=0
				dtect_tag=''
			elif waitforelse==1:
				start=0
				dtect_tag=''
			else:
				# print('Tag_detected: ',dtect_tag)
				# print('Tag Text: ',tag_text)
				if z =='FALSE':
					waitforelse=1
				else:
					text_replace.append(tag_text)
				z=evaluateTag(list_df,dtect_tag,i)
				# print('eval tag : ',z)
				if_tags_order.append(z)
				else_count=else_count+1
				# print('text Replace :',text_replace)
				# print('if_tags_order:',if_tags_order)			
				tag_text=''
				dtect_tag=''
				start=0
		if dtect_tag in other_tags:
			# print('Tag_detected other: ',dtect_tag)
			# print('Tag Text: ',tag_text)
			if waitforendif==1:
				if dtect_tag==other_tags[1]:
					# print('ENDIF')
					start=0
					dtect_tag=''
					tag_text=''
					waitforendif=0
			else:
				if z=='TRUE':
					text_replace.append(tag_text)
				# print('text Replace :',text_replace)
				if dtect_tag==other_tags[0]:
					tag_eval=if_tags_order[else_count-1]
					else_count=else_count-1
					# print('tag Eval :',tag_eval)
					if tag_eval=='TRUE':
						waitforendif=1
						start=0
						dtect_tag=''
						tag_text=''
					else:
						z='TRUE'
						dtect_tag=''
						tag_text=''
						start=0
				else:
					dtect_tag=''
					tag_text=''
					start=0
					z='TRUE'

		# print(dtect_tag)
		# print(other_tags)
		if dtect_tag in leave_tags:
			# print('dtect_tag extra:',dtect_tag)
			tag_text=tag_text+dtect_tag
			start =0
			dtect_tag=''
	# print(text_replace)
	return (''.join(text_replace))


def find_and_replace_if_tags(list_df,document,i,target_file):
	"""
	Performs processing for IF-ELSE-ENDIF tags and FILE tags.  Produces
	a data structure reflecting the content of an output Word file prior
	to replacing TEXT and IMAGE tags, if any, within that file
	:param list_df: (list of Pandas dataframes) output from Pandas from reading Excel workbook
	:param document: (python-docx.Document) Document instance
	:param i: (integer) current iterator value (0...len-1)
	:return: (python-docx.Document) a modified Document instance with
	IF tags replaced with conditional content
	"""
	# tags_associated=['[[ELSE]]',]

	# Iterate through paragraph and table entities within document
	# through iter_block_items(), which emits python-docx Paragraph
	# and Table instances for such entities
	for block in iter_block_items(document):

		if isinstance(block, Paragraph):
			# current block is an instance of python-docx.Paragraph
			block_text=block.text # retrieve unformatted text string
			if '[[IF:' in block_text: # detect start of IF tag
				# print('btt',block_text)

				## ANALYSIS NOT ATTEMPTED
				z=replaceIftext(list_df,block_text,i)
				block.text=z
			# print('bt ',block_text)
		else:
			print("Current block is not a paragraph")

	document.save(target_file)
	return document


def Document_data(document):
	"""
	Returns a string containing all unformatted content from
	all Paragraph instances in the provided Document instance
	:param document: (python-docx.Document) input Word file
	:return: (string) unformatted content from all Paragraph
	instances within the provided Document instance
	"""
	Run_list=[]

	# Create empty content string prior to appending content
	# to that string below

	# Iterate through each Paragraph and Table instance
	# within the provided Document
	for block in iter_block_items(document):
		# Only process content for Paragraph instances
		if isinstance(block, Paragraph):
			# Retrieve unformatted text from Paragraph
			block_text=block.text

			# print('bt ',block_text)
			for run in block.runs:
				Run_list.append(run.text)

	# Return string containing all unformatted text from
	# block (which can only be a Paragraph)
	return Run_list

def replace_file_tags(list_df,document,files_dict,index):
	"""
	Replaces ONE [[FILE...]] tag in python-docx.Document instance with
	content from corresponding Word file.  Caller will repeatedly call
	this function until instance is false, reflecting that no [[FILE...]]
	tags were processed, to ensure that all [[FILE...]] tags within the
	Document instance have been processed
	:param list_df: (list of Panda data frames) worksheets
	:param document: (python-docx.Document instance) Word document
	whose FILE tags will be replaced
	:param files_dict: (dictionary) keys as filepaths and values as
	string of unformatted content from Paragraphs in the Word document
	at that filepath
	:param index: (int) 0..n-1 loop index
	:return: document,instance: (python-docx.Document, bool) document
	is the modified Word file after replacing [[FILE...]] tags and
	instance is a flag indicating whether any [[FILE...]] tags were
	replaced
	"""

	# instance is True if [[FILE...]] tag replacement has occurred
	instance=False

	# Iterate through each python-docx.Paragraph or Table instance
	# in received Document instance
	for block in iter_block_items(document):

		# Process as python-docx.Paragraph instance if it is one
		if isinstance(block, Paragraph):
			# Retrieve unformatted text from current block
			block_text=block.text
			# print('bt ',block_text)

			# Retrieve FILES data frame for FILES workbook
			file_df=list_df[3]

			# Check whether start of FILE tag is present within
			# unformatted text for block.  Unformatted text is a
			# viable way to check for the tag because we don't
			# care about formatting considerations (distinguishing
			# Runs) for this check
			if '[[FILE:' in block_text:

				# Parse unformatted text into substrings on the
				# [[FILE...]] tag start.  Produces a list of substrings
				# with the 0th element reflecting text preceding the
				# initial [[FILE...]] tag
				new_p = OxmlElement("w:p")
				block._p.addnext(new_p)
				new_para = Paragraph(new_p, block._parent)
				for run in block.runs:
					if '[[FILE:' in run.text:
						substrings=(run.text).split('[[FILE:')
						new_para.add_run(substrings[0])

						# Iterate through all substrings (excluding the initial
						# substring because it precedes the first tag)
						for substring in substrings[1:]:

							# Extract the [[FILE...]] tag's identifier
							# (which indexes a filepath in the data frame)
							# from the current substring

							identifier=substring.split(']]')[0]
							if '(' in identifier:
								identifier = identifier.split('(')[0]
								comment=identifier.split('(')[1].split(')')[0]
								tag='[[FILE:'+identifier+'('+comment+')'+']]'
							else:
								tag = '[[FILE:' + identifier + ']]'
							text_add=substring.replace((identifier+']]'),'')
							# print('idd',identifier)

							# Recreate full FILE tag from identifier


							# print('filename',file_df[identifier])

							# Retrieve filepath for current identifier and current
							# index value from files data frame (FILES workbook
							# data)
							try:
								filepath = list(file_df[identifier])[index]
								List_runs=files_dict[filepath]
								for tag_run in List_runs:
									new_para.add_run(tag_run)
								new_para.add_run(text_add)
							except Exception as e:
								new_para.add_run(tag)
								new_para.add_run(text_add)
								print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))
							# 	Need to add a good way for validation. This can be a corner case.
							# print('block_text',block_text)
							# block_text=block.text
					else:
						new_para.add_run(run.text)
				delete_paragraph(block)
				instance=True
				# Store unformatted text for current Paragraph instance after
				# replacing [[FILE...]] tags with corresponding file content
				# (although see Question 7 above)

				# Set instance flag reflecting that [[FILE...]] tag
				# replacement has occurred

		# Process as python-docx.Table instance if it is one
		elif isinstance(block, Table):
			# print('Enter Table: ',table)
			# ll

			# Iterate through each python-docx.Row instance within
			# current python-docx.Table instance
			for row in block.rows:

				# Iterate through each python-docx.Cell instance within
				# current python-docx.Row instance
				for cell in row.cells:

					# Iterate through each python-docx.Paragraph instance
					# within current python-docx.Cell instance
					for paragraph in cell.paragraphs:

						# Extract unformatted text from current Paragraph instance
						block_text=paragraph.text

						# print('bt ',block_text)

						# Retrieve Pandas data frame for FILES worksheet
						file_df=list_df[3]

						# Check whether start of FILE tag is present within
						# unformatted text for block.  Unformatted text is a
						# viable way to check for the tag because we don't
						# care about formatting considerations (distinguishing
						# Runs) for this check
						if '[[FILE:' in block_text:
							new_p = OxmlElement("w:p")
							paragraph._p.addnext(new_p)
							new_para = Paragraph(new_p, paragraph._parent)
							for run in paragraph.runs:
								if '[[FILE:' in run.text:
									substrings = (run.text).split('[[FILE:')
									new_para.add_run(substrings[0])

									# Iterate through all substrings (excluding the initial
									# substring because it precedes the first tag)
									for substring in substrings[1:]:

										# Extract the [[FILE...]] tag's identifier
										# (which indexes a filepath in the data frame)
										# from the current substring

										identifier = substring.split(']]')[0]
										if '(' in identifier:
											identifier = identifier.split('(')[0]
											comment = identifier.split('(')[1].split(')')[0]
											tag = '[[FILE:' + identifier + '(' + comment + ')' + ']]'
										else:
											tag = '[[FILE:' + identifier + ']]'
										text_add = substring.replace((identifier + ']]'), '')
										# print('idd',identifier)

										# Recreate full FILE tag from identifier

										# print('filename',file_df[identifier])

										# Retrieve filepath for current identifier and current
										# index value from files data frame (FILES workbook
										# data)
										try:
											filepath = list(file_df[identifier])[index]
											List_runs = files_dict[filepath]
											for tag_run in List_runs:
												new_para.add_run(tag_run)
											new_para.add_run(text_add)
										except Exception as e:
											new_para.add_run(tag)
											new_para.add_run(text_add)
											print('Error the Tag {} is not present in the Excel Sheet.'.format(tag))
								# print('block_text',block_text)
								# block_text=block.text
								else:
									new_para.add_run(run.text)
							# Set instance flag reflecting that [[FILE...]] tag
							# replacement has occurred
							instance=True
							delete_paragraph(paragraph)
							# print('Replaced: ',paragraph.text)

		else:
			print("Warning: current block is neither a Paragraph nor a Table instance!")

	# Return python-docx.Document instance with zero or one [[FILE...]] tag replaced
	# with content from corresponding Word file (if replacement occurred) and instance
	# flag as True if replacement occurred
	return document,instance

def load_file_dict(filepaths):
	"""
	Creates a dictionary of filepath-content_string pairs from
	keys for each filepath in the list of provided filepaths and
	values as a string comprising all unformatted content from
	Paragraphs only within documents at the provided filepaths
	:param filepaths: (list of string) filepaths for all FILE
	tags across all iterator values
	:return: (dict of string-string pairs) filepath keys and
	unformatted string of Paragraph content from filepath values
	"""
	# Create an empty dictionary for receiving file data
	files_dict={}

	# Iterate through each filepath received as ARGV
	for filepath in filepaths:
		# Create a python-docx.Document instance for current
		# filepath and read corresponding Word file into instance
		document=Document(filepath)

		# Retrieves all unformatted content from Paragraphs in the
		# Document instance as values in a dictionary that are keyed
		# from corresponding filepaths
		files_dict[filepath]=Document_data(document)

	# print(files_dict)
	return files_dict

def preprocess_files(input_wordfile,list_df):
	"""
	Begins by conditionally including or excluding content, based on
	resolving GLOBAL values for [[IF...]] tags.  Continues by replacing
	[[FILE...]] tags with content from corresponding Word files.  Returns
	a Document instance reflecting such changes
	:param input_wordfile: (string) filepath for input Word file
	:param list_df: (list of Pandas data frames) workbook data
	:return: (python-docx.Document) input Word file after conditionally
	including or excluding content based on [[IF...]] tags and after
	replacing [[FILE...]] tags with content from corresponding Word files
	"""

	# Generate a list of all unique filepaths for all FILE tags
	# for all iterator values.  This list will be used for parsing
	# all files upfront, and only once, notwithstanding that those
	# files could be included in several ways and in several
	# locations in other files (efficient Word file parsing)
	filepaths=getallfiles(list_df)

	# Generates a dictionary with filepaths as keys and
	# strings containing all unformatted content from Paragraphs
	# within documents at the corresponding filepath as values.
	files_dict=load_file_dict(filepaths)

	# Check that input Word file is not in list of filepaths because
	# that situation would reflect a circular FILE tag situation
	if input_wordfile in filepaths:
		print("ERROR: input Word file is also a filepath for a FILE tag - circular FILE problem")
		exit(1)

	# Add input Word file to end of list of unique filepaths for FILE tags
	# for all iterator values
	filepaths.append(input_wordfile)

	# Iterate through each filepath in list of filepaths for all FILE
	# tags across all iterators (plus input Word file)
	for index in range(len(filepaths)):

		# Retrieve current file from list of files
		filepath=filepaths[index]

		# Determine if current file is last file in list (would
		# be the input Word file) because input Word file was added
		# to the end of that list
		if index==len(filepaths)-1:
			# Generate a dictionary of key-value pairs with keys as
			# filepaths (Word files identified in [[FILE...]] tags
			# and input Word file) and strings of unformatted content
			# extracted from Paragraphs within filepath
			files_dict=load_file_dict(filepaths)
			# ll
		document=Document(filepath)
		# print("WARNING: skipping IF tag parsing and replacement")
		find_and_replace_if_tags(list_df,document,index,filepath)
		# ll

		# Read Word file at filepath into a python-docx.Document
		document=Document(filepath)
		# content=Document_data(document)

		# Repeatedly call replace_file_tags() until that function returns false,
		# indicating that all [[FILE...]] tags have been replaced
		while True:
			# Replace [[FILE...]] tags within document with content from
			# corresponding Word files
			document,tag_replacement_performed=replace_file_tags(list_df,document,files_dict,index)

			# print('contain: ',tag_replacement_performed)

			# Determine whether looping should continue based on flag
			# returned from replace_file_tags()
			if tag_replacement_performed==True:
				pass
			else:
				break	
#  This code your wrote is not correct  since it returns the first instance of the document. what the idea here is to save each file we open we save it so that the next time we read it 
#  we hve the updted one. If we return from here then the function won't get completed.
		# # [[FILE...]] tag replacement is complete so return resulting Document instance
		# return document
		document.save(filepath)
	# return document

def validatesheetnames(sheet_names):
	if sheet_names[0]=='GLOBALS' and sheet_names[1]=='TEXT' and sheet_names[2]=='IMAGE' and sheet_names[3]=='FILE':
		return 1
	return 0

def replacehead(df):
	new_header = df.iloc[0] #grab the first row for the header
	df = df[1:] #take the data less the header row
	df.columns = new_header #set the header row as the df header
	return df


def cleanExcel(df):
	df_globals=replacehead(df.parse('GLOBALS').dropna(how='all').dropna(axis='columns'))
	df_text=replacehead(df.parse('TEXT').dropna(how='all').dropna(axis='columns'))
	df_image=replacehead(df.parse('IMAGE').dropna(how='all').dropna(axis='columns'))
	df_file=replacehead(df.parse('FILE').dropna(how='all').dropna(axis='columns'))
	# print(df_file)
	return df_globals,df_text,df_image,df_file

def read_excel(file_name):
	"""
	Read Excel workbook, whose filepath is specified as file_name, into Pandas
	and return a list of Pandas dataframes - one each for globals, text tags,
	image tags and file tags or return 0 if Excel workbook not parsed properly
	:param file_name: (string) filepath for Excel workbook
	:return: (list of dataframes or 0 if error) list of dataframes for globals,
	text tags, image tags and file tags
	"""
	df=pd.ExcelFile(file_name)
	# print(df.sheet_names)
	if validatesheetnames(df.sheet_names):
		df_globals,df_text,df_image,df_file=cleanExcel(df)
		# print(df_globals)
		# print(df_text)
		# print(df_image)
		# print(df_file)
		return [df_globals,df_text,df_image,df_file]
	else:
		return 0	

def validate_excel():
	"""
	TO DO 
	"""

def replace_text_tags(list_df,document,index,target_file):
	"""
	Replaces [[TEXT:<identifier>]] tags with text corresponding to the identifier
	as specified in the TEXT worksheet and for the current index.  Uses the
	"text" attribute of a Paragraph instance to retrieve the unformatted text
	for the paragraph.  We are able to use the unformatted text version of the
	paragraph because we require that formatting is consistent for a tag, which
	results in tags that do not span a "Run" in python-docx parlance
	:param list_df: list of Pandas dataframes reflecting worksheets
	for globals, text tags, image tags and file tags
	:param document: (python-docx.Document) Document instance
	:param index: (integer) current index value
	:return: (python-docx.Document) modified Document instance
	"""

	# Iterate through each Paragraph and Table within Document
	# based on iter_block_items(), which emits Paragraph and
	# Table instances within the Document
	for block in iter_block_items(document):

		# Process block as paragraph if current block is an instance of Paragraph
		if isinstance(block, Paragraph):

			# Extract unformatted text from Paragraph instance
			block_text=block.text
			# print(block_text)

			# Determine whether [[TEXT:<identifier>(optional-comment)]] is present
			if '[[TEXT:' in block_text:

				# At least one [[TEXT...]] tag is present so find substrings for
				# the one or more [[TEXT...]] tags in the current paragraph.  Returns
				# a list of substrings separated on '[[TEXT:'' such that the initial
				# element is the string preceding the first [[TEXT...]] tag, the
				# next element is text after [[TEXT: until the start of the next
				# [[TEXT...]] tag or the end of line, repeating thereafter.  Thus,
				# the first characters in the first element are the identifier and
				# optional comment for the [[TEXT...]] tag
				new_p = OxmlElement("w:p")
				block._p.addnext(new_p)
				new_para = Paragraph(new_p, block._parent)
				for run in block.runs:
					if '[[TEXT:' in run.text:
						numtexttags=(run.text).split('[[TEXT:')
						new_para.add_run(numtexttags[0])
						# Iterate through each [[TEXT...]] tag, after skipping text preceding
						# the first [[TEXT...]] tag (the 1: expression)
						for tag in numtexttags[1:]:
							# for current substring, split the end of the tag from text
							# after the end of the tag.  The identifier is to the left of
							# the initial substring
							identifier=tag.split(']]')[0]
							if '(' in identifier:
								identifier = identifier.split('(')[0]
								comment=identifier.split('(')[1].split(')')[0]
								tag_d='[[TEXT:'+identifier+'('+comment+')'+']]'
							else:
								tag_d = '[[TEXT:' + identifier + ']]'
							text_add=tag.replace((identifier+']]'),'')
							# Recreate tag so it can be replaced easily in Paragraph instance
							# tag_d='[[TEXT:'+identifier+']]' # Question2: will this later include comment parsing?
							#  try catch Required to avoid key error in pandas and is used to validate and detect invalid tag.If you want to change it we can discuss
							try:
								# print('Tag ID : ',tag_d)
								text_rep=list(list_df[1][identifier])[index]
								new_para.add_run(text_rep)
								new_para.add_run(text_add)
								# Replace [[TEXT...]] tag with its corresponding text
								# from TEXT worksheet
								# block_text=block_text.replace(tag_d,text_rep)
								# print('block_text: ',block_text)
							except Exception as e:
								new_para.add_run(tag_d)
								new_para.add_run(text_add)
								print('Error the Tag {} is not present in the Excel Sheet.'.format(tag_d))
					else:
						new_para.add_run(run.text)
				delete_paragraph(block)
			else:
				# No [[TEXT...]] tag present in paragraph
				pass
		# Process block as table if current block is an instance of Table
		elif isinstance(block, Table):
			# Iterate through each Row instance within the Table
			for row in block.rows:
				# Iterate through each Cell instance within the Row
				for cell in row.cells:
					# Iterate through each Paragraph instance within the Cell
					for paragraph in cell.paragraphs:
						# Extract unformatted text from Paragraph instance
						block_text=paragraph.text

						# NOTE THAT THIS CODE REPEATS FUNCTIONALITY FOR PARAGRAPHS

						# Future project: break logic below into function that is
						# called for Paragraph instances and Table instances

						# Detect presence of [[TEXT...]] tag in current block
						if '[[TEXT:' in block_text:
							# [[TEXT...]] tag found so parse it

							# Create a list of substrings, the first of which
							# ends just before the first [[TEXT:...]] tag, and
							# the remainder of which begins at the start of each
							# [[TEXT...]] tag
							new_p = OxmlElement("w:p")
							paragraph._p.addnext(new_p)
							new_para = Paragraph(new_p, paragraph._parent)
							for run in paragraph.runs:
								if '[[TEXT:' in run.text:
									numtexttags = (run.text).split('[[TEXT:')
									new_para.add_run(numtexttags[0])
									# Iterate through each [[TEXT...]] tag, after skipping text preceding
									# the first [[TEXT...]] tag (the 1: expression)
									for tag in numtexttags[1:]:
										# for current substring, split the end of the tag from text
										# after the end of the tag.  The identifier is to the left of
										# the initial substring
										identifier = tag.split(']]')[0]
										if '(' in identifier:
											identifier = identifier.split('(')[0]
											comment = identifier.split('(')[1].split(')')[0]
											tag_d = '[[TEXT:' + identifier + '(' + comment + ')' + ']]'
										else:
											tag_d = '[[TEXT:' + identifier + ']]'
										text_add = tag.replace((identifier + ']]'), '')
										# Recreate tag so it can be replaced easily in Paragraph instance
										# tag_d='[[TEXT:'+identifier+']]' # Question2: will this later include comment parsing?
										#  try catch Required to avoid key error in pandas and is used to validate and detect invalid tag.If you want to change it we can discuss
										try:
											# print('Tag ID : ',tag_d)
											text_rep = list(list_df[1][identifier])[index]
											new_para.add_run(text_rep)
											new_para.add_run(text_add)
										# Replace [[TEXT...]] tag with its corresponding text
										# from TEXT worksheet
										# block_text=block_text.replace(tag_d,text_rep)
										# print('block_text: ',block_text)
										except Exception as e:
											print('Error the Tag {} is not present in the Excel Sheet.'.format(tag_d))
								else:
									new_para.add_run(run.text)
							delete_paragraph(paragraph)
		else:
			print("Warning: Current block is neither a Paragraph nor a Table")

	# Return Document instance to caller so next processings step can build on
	# this result
	document.save(target_file)
	return document


def insert_image_after(run,image,img_width=-1,img_height=-1):
	"""
	Inserts the provided image into the provided run, optionally with
	width and height dimensions, and directs python-docx to scale based
	on width and/or height dimensions
	:param run: (python-docx.Run) Run instance into which image is inserted
	:param image: (string) filepath of image to insert
	:param img_width: (float) width of image in inches
	:param img_height: (float) height of image in inches
	:return: (empty)
	"""

	if img_height==-1:
		# image height not specified so check image width
		if img_width==-1:
			# image height and width not specified so let python-docx
			# determine image dimensions
			run.add_picture(image)
		else:
			# image width but not height specified so tell python-docx
			# to scale the image proportionally for the width
			run.add_picture(image,Inches(img_width))
	else:
		# image height specified so check image width
		if img_width==-1:
			# image height specified but not image width so tell
			# python-docx to scale image proportionally for the height
			run.add_picture(image, Inches(img_height))
		else:
			# image width and height provided so tell python-docx
			# to scale the image as provided in both dimensions
			run.add_picture(image, width=Inches(img_width), height = Inches(img_height))
    
def insert_run_after(paragraph, text, style=None):
	"""
	Uses internal python-docx methods to add a new Run instance, containing the
	provided text, at the end of the provided Paragraph instance
	
	See https://stackoverflow.com/questions/48663788/python-docx-insert-a-paragraph-after
	:param paragraph: (python-docx.Paragraph) a paragraph
	:param text: (string) text to add
	:param style: (not currently implemented) see https://python-docx.readthedocs.io/en/latest/api/style.html
	:return: (empty)
	"""

	# Create a new data structure for receiving text and adding data
	# structure to received paragraph
	new_p = OxmlElement("w:p")

	# Add new data structure to received paragraph
	paragraph._p.addnext(new_p)

	# Create new Paragraph instance linked to parent Paragraph instance
	# using new data structure
	new_para = Paragraph(new_p, paragraph._parent,style)

	# Add text to new Paragraph instance as a new Run instance
	new_para.add_run(text)
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def replace_image_tag(list_df,document,index,target_file):
	"""
	Replaces [[IMAGE:<identifier>]] tags with an image corresponding to the identifier
	as specified in the IMAGES worksheet and for the current index.  Uses the
	"text" attribute of a Run instance to retrieve the unformatted text
	for the paragraph.  We are able to use the unformatted text version of the
	paragraph because we require that formatting is consistent for a tag, which
	results in tags that do not span a "Run" in python-docx parlance
	:param list_df: list of Pandas dataframes reflecting worksheets
	for globals, text tags, image tags and file tags
	:param document: (python-docx.Document) Document instance
	:param index: (integer) current index value
	:return: (python-docx.Document) modified Document instance
	"""

	# Iterate through each Paragraph and Table within Document
	# based on iter_block_items(), which emits Paragraph and
	# Table instances within the Document
	for block in iter_block_items(document):

		# Process block as paragraph if current block is an instance of Paragraph
		if isinstance(block, Paragraph):
			# Extract unformatted text from Paragraph instance
			block_text=block.text
			
			if '[[IMAGE:' in block.text:
				# Iterate through each Run within current block
				# (a Paragraph instance)
				new_p = OxmlElement("w:p")
				block._p.addnext(new_p)
				new_para = Paragraph(new_p, block._parent)
				for run in block.runs:
					# print('run : ',run.text)
					try:
						# Determine whether [[IMAGE:<identifier>(optional-comment)]] is present
						if '[[IMAGE:' in run.text:
							# At least one [[IMAGE...]] tag is present so find substrings for
							# the one or more [[IMAGE...]] tags in the current run.  Returns
							# a list of substrings separated on '[[IMAGE:'' such that the 0th
							# element is the string preceding the first [[IMAGE...]] tag, the
							# 1st element is text after [[IMAGE: until the start of the next
							# [[IMAGE...]] tag or the end of line, repeating thereafter.  Thus,
							# the first characters in the first element are the identifier and
							# optional comment for the [[IMAGE...]] tag

							image_items=(run.text).split('[[IMAGE:')
							# print('image_items: ',image_items)
							new_para.add_run(image_items[0])
							for image in image_items[1:]:
								# print('split image : ',image.split(']]'))
								image_tag='[[IMAGE:'+image.split(']]')[0]+']]'
								identifier=image.split(']]')[0]
								if '(' in identifier:
									identifier = identifier.split('(')[0]
									comment = identifier.split('(')[1].split(')')[0]
									tag_d = '[[TEXT:' + identifier + '(' + comment + ')' + ']]'
								else:
									tag_d = '[[TEXT:' + identifier + ']]'
								text_add=image.replace((identifier+']]'),' ')
								# print('text_add: ' ,text_add)
								try:
									img_path=list(list_df[2][identifier])[index]
									img_width=list(list_df[2][identifier+'_width'])[index]
									img_height=list(list_df[2][identifier+'_height'])[index]
									insert_image_after(new_para.add_run(),img_path,img_width,img_height)
									new_para.add_run(text_add)
								except Exception as e:
									print(e)
									new_para.add_run(tag_d)
									new_para.add_run(text_add)
									# block.text=(block.text).replace(text_add,'')
									# print('ErrorPar: ',str(e))
									print('Error the Tag {} is not present in the Excel Sheet.'.format(image_tag))
						else:
							new_para.add_run(run.text)
					except Exception as e:
						# print('ErrorPar: ',str(e))
						print('Image Tag does not fall under  Run.')
						# run.text=''
				delete_paragraph(block)
		# Process block as table if current block is an instance of Table
		elif isinstance(block, Table):
			# Iterate through each Row instance within the Table
			for row in block.rows:
				# Iterate through each Cell instance within the Row
				for cell in row.cells:
					# Iterate through each Paragraph instance within the Cell
					for paragraph in cell.paragraphs:
						# Extract unformatted text from Paragraph instance
						block_text=paragraph.text

						if '[[IMAGE:' in paragraph.text:
							new_p = OxmlElement("w:p")
							paragraph._p.addnext(new_p)
							new_para = Paragraph(new_p, paragraph._parent)
							for run in paragraph.runs:
								# print('run: ',run.text)
								try:
									if '[[IMAGE:' in run.text:
										image_items=(run.text).split('[[IMAGE:')
										# print('image_items: ',image_items)
										new_para.add_run(image_items[0])
										for image in image_items[1:]:
											# print('split image : ',image.split(']]'))
											image_tag='[[IMAGE:'+image.split(']]')[0]+']]'
											identifier=image.split(']]')[0]
											if '(' in identifier:
												identifier = identifier.split('(')[0]
												comment = identifier.split('(')[1].split(')')[0]
												# text_add = '[[TEXT:' + identifier + '(' + comment + ')' + ']]'
											# else:
												# text_add = '[[TEXT:' + identifier + ']]'
											text_add = image.replace((identifier + ']]'), ' ')

											# print('text_add: ' ,text_add)
											try:
												img_path=list(list_df[2][identifier])[index]
												img_width=list(list_df[2][identifier+'_width'])[index]
												img_height=list(list_df[2][identifier+'_height'])[index]
												insert_image_after(new_para.add_run(),img_path,img_width,img_height)
												new_para.add_run(text_add)
											except Exception as e:
												print(e)
												new_para.add_run(text_add)
												# block.text=(block.text).replace(text_add,'')
												# print('ErrorPar: ',str(e))
												print('Error the Tag {} is not present in the Excel Sheet.'.format(image_tag))
									else:
										new_para.add_run(run.text)
								except Exception as e:							
									# print('ErrorTable: ',str(e))
									print('Image Tag Doesnt fall under run ')

							delete_paragraph(paragraph)
	document.save(target_file)
	return document # with modifications

def process_document(list_df,preprocessed_document,index,target_file,input_file):
	"""
	Receives a Document instance that has been preprocessed (all [[FILE...]]
	and [[IF...]] tags have been replaced) and processes that Document
	instance by replacing [[IMAGE...]] and [[TEXT...]] tags with their
	corresponding content
	:param list_df: (list of Pandas dataframes) comprises data frames
	for globals, text tags, image tags and file tags from spreadsheet
	:param preprocessed_document: (python-docx.Document) Document
	instance storing contents of input Word file that will be modified
	but which also has [[FILE...]] and [[IF...]] tags already processed
	:param index: (int) current index value (determines which row of
	Excel data will be used)
	:return: (python-docx.Document) Document instance read from
	input Word file after parsing for all IF/FILE/IMAGE/TEXT tags
	"""

	# Parse Document instance for IF-ELSE-ENDIF tags and replace those
	# tags with appropriate conditional content

	print("Warning: IF tag processing is commented out")
	# NOTE - commenting this out because [[IF...]] tag processing
	# is not ready
	document = find_and_replace_if_tags(list_df,preprocessed_document,index,target_file)
	preprocessed_document=Document(target_file)

	# Parse Document instance for TEXT tags and replace those
	# tags with appropriate text
	document_after_text_tag_replacement = replace_text_tags(list_df,preprocessed_document,index,target_file)
	preprocessed_document=Document(target_file)
	# Parse Document instance for IMAGE tags and replace those
	# tags with appropriate image content
	document_after_image_tag_replacement = replace_image_tag(list_df,preprocessed_document,index,target_file)

	# return document_after_image_tag_replacement
	# return document_after_text_tag_replacement

if __name__ == '__main__':

	print('Loading And Validating EXCEL')

	# Create filepath for Excel workbook as CWD/Files/<workbook>
	excel_file_path=os.path.join(os.path.join(os.getcwd(),'Files'),'excelfile.xlsx')

	# Return a list of Pandas dataframes for globals, text tags, image tags and file tags
	# from separate worksheets (for each) within Excel workbook
	list_df=read_excel(excel_file_path)

	# Handle error from Pandas processing of Excel workbook
	if type(list_df) is int:
		print("Error reading Excel workbook at %s.  Exiting." % excel_file_path)
		exit(1)

	# Create filepath for input document as CWD/Files/<input-file>
	input_wordfile=os.path.join(os.path.join(os.getcwd(),'Files'),'inputDocument.docx')
	
	## NOT USED
	# generate_tags_list(input_wordfile,list_df)

	# Generate a list of Word files to be generated - one per iterator value -
	# from the GLOBALS worksheet
	target_files=list(list_df[0]['DESTINATION'])
	# print(target_file)
	# ll

	## Iterate through each Word file to be generated
	for index in range(len(target_files)):
	
		print('Preprocessing Files')
		# Processes [[IF...]] and [[FILE...]] tags to produce Document
		# instances reflecting the result of that processing (although
		# [[IF...]] tag processing is a work in progress so I commented
		# it out
		preprocessed_document = preprocess_files(input_wordfile,list_df)
		
		print('Iteration ',index+1)
		## Retrieve name of current Word file to be created
		file_target = target_files[index]
		
		# Find and operate on [[IMAGE...]] and [[TEXT...] tags to
		# produce a modified Document instance
		preprocessed_document=Document(input_wordfile)
		processed_document = process_document(list_df,preprocessed_document,index,file_target,input_wordfile)


		# Now save the result as the target Word file
		# processed_document.save(file_target)

	# read_document(list_df,input_wordfile,target_file[0],0)